import os
import re
import io
import time
import uuid
import json
import copy
import base64
import shutil
import smtplib
import pathlib
import traceback
import subprocess
import datetime
import mimetypes
from io import BytesIO
from decimal import Decimal, getcontext
from concurrent.futures import ThreadPoolExecutor, as_completed
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication

import fitz                     
import requests
import mysql.connector
from mysql.connector import pooling
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from PyPDF2 import PdfReader
from jinja2 import Template
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from flask import Flask, request, jsonify, send_file, send_from_directory, abort
from flask_cors import CORS
from mistralai import Mistral

app = Flask(__name__)
CORS(app)
load_dotenv()

db_config = {
    'host': '88.202.183.189',
    'user': 'petatron_usr',
    'password': 'Qxg655hwyQyThYLW',
    'database': 'petatron_aijobhunter_dev'
}

connection_pool = pooling.MySQLConnectionPool(
    pool_name="mypool",
    pool_size=10,
    pool_reset_session=True,
    **db_config
)

MISTRAL_API_KEY = os.getenv("MISTRAL_API_KEY")
if not MISTRAL_API_KEY:
    raise RuntimeError("MISTRAL_API_KEY is missing. Set it in your environment or .env")

MODEL = os.getenv("MISTRAL_MODEL", "pixtral-12b-2409")
client = Mistral(api_key=MISTRAL_API_KEY)
RESUME_FOLDER = "generated_resumes"
os.makedirs(RESUME_FOLDER, exist_ok=True)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_FOLDER = os.path.join(BASE_DIR, "generated_resumes")
TEMPLATE_PATH = os.path.join(TEMPLATE_FOLDER, "classic_template_aligned.docx")
mistral_client = Mistral(api_key=MISTRAL_API_KEY)
OUTPUT_DIR = "generated_resumes"
os.makedirs(OUTPUT_DIR, exist_ok=True)

#INSERT AI LOG FUNCTION
def insert_ai_log(conn, job_id, user_id, activity_type, service_name, model_name,
                  user_input, ai_response_text, response_time_ms, status,
                  error_message, metadata, created_at=None):
    if created_at is None:
        created_at = int(time.time())
    cursor = conn.cursor()
    insert_query = """
        INSERT INTO tbl_users_ai_activities 
        (job_id, user_id, activity_type, service_name, model_name, user_input, 
         ai_response, response_time_ms, status, error_message, metadata, created_at)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
    """
    data = (job_id, user_id, activity_type, service_name, model_name, user_input,
            ai_response_text, response_time_ms, status, error_message, metadata, created_at)
    cursor.execute(insert_query, data)
    conn.commit()
    cursor.close()

# File extraction utilities
def extract_text_from_pdf(path):
    doc = fitz.open(path)
    try:
        return "".join(page.get_text() for page in doc).strip()
    finally:
        doc.close()

def extract_text_from_docx(path):
    doc = Document(path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    return "\n".join(full_text)

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported format (only .pdf and .docx supported)")

def extract_links_from_pdf(path):
    doc = fitz.open(path)
    try:
        out = []
        for page in doc:
            for link in page.get_links():
                uri = link.get("uri")
                if uri:
                    out.append(uri)
        return out
    finally:
        doc.close()

def extract_links_from_docx(path):
    doc = Document(path)
    return [rel.target_ref for rel in doc.part.rels.values() if "hyperlink" in rel.reltype]

def extract_urls_from_text(text):
    url_pattern = r"(https?://[^\s)>\]]+|www\.[^\s)>\]]+|linkedin\.com[^\s]*|github\.com[^\s]*)"
    return re.findall(url_pattern, text, re.IGNORECASE)

def extract_social_links(links):
    def first_contains(substr):
        for l in links:
            if substr in l.lower():
                return l
        return ""
    return {
        "linkedin": first_contains("linkedin.com"),
        "github": first_contains("github.com"),
    }

def clean_response_json(response_text):
    match = re.search(r"```json\s*(\{.*?\})\s*```", response_text, re.DOTALL)
    if not match:
        match = re.search(r"(\{[\s\S]*\})", response_text)
    if not match:
        return {"error": "Could not parse JSON"}
    json_str = match.group(1)
    json_str = re.sub(r",\s*(\}|\])", r"\1", json_str)
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        return {"error": f"JSON decoding error: {str(e)}"}

def _truncate(text: str, max_chars: int = 12000) -> str:
    if not text:
        return ""
    return text[:max_chars]

def _normalize_yes_no(s: str) -> str:
    if not isinstance(s, str):
        return "no"
    s = s.strip().lower()
    if s.startswith("yes"):
        return "yes"
    if s.startswith("no"):
        return "no"
    return "no"

def is_resume_with_mistral(doc_text: str) -> str:
    """
    Returns strictly "yes" or "no" (lowercase).
    """
    prompt = f"""
You are a binary classifier. Decide if the following document text is a resume/CV (a personal job-seeking document listing a person's education, skills, projects, experience).

Return exactly one word: "yes" or "no". No punctuation. No explanation.

Document text (may be partial):
\"\"\" 
{_truncate(doc_text)}
\"\"\" 
"""
    resp = client.chat.complete(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
    )
    answer = resp.choices[0].message.content.strip()
    return _normalize_yes_no(answer)

def extract_info_with_mistral(resume_text: str, transcripts: str = "") -> str:
    transcript_block = f"""
Video Transcripts:
\"\"\"
{transcripts}
\"\"\"
""" if transcripts.strip() else ""
 
    prompt = f"""
You are an expert resume parser. Extract information from the candidate’s **resume text** and **transcript block**
If information is present in transcripts but missing in the resume(except employement) include it in the JSON schema.
CRITICAL RULES (read carefully; violating any makes the output INVALID):
- Merge resume and transcript information into one JSON.
- If a detail appears only in transcripts, include it.
- If the same detail appears in both, prefer the resume version.
- Do not hallucinate. If info not found, use "" or [] exactly.
- Keep the schema exactly as shown. Do not add, remove, rename, or reorder fields.
- **ROLE CLASSIFICATION**:
  • If the title contains any of: ["intern", "internship", "trainee"] (case-insensitive) → put this job ONLY under "internship".
  • All other paid roles go under "EMPLOYMENT".
  • Never place an intern/trainee role in "EMPLOYMENT".
- Normalize any “Courses/Online Courses/MOOCs/Training Programs” into the "CERTIFICATION" array.
- "course_type" allowed values: "Full-time" | "Part-time" | "Distance" | "" (empty string if not explicitly mentioned).
- Normalize any “Research projects, Publications, Journal Papers, Conference Papers” into the "ACHIEVMENT" array.
 
Resume Text:
\"\"\"
{resume_text}
\"\"\"
Transcript:
{transcript_block}
 
Return strictly valid JSON in this schema:
{{
  "first_name": "",    // first name of the candidate
  "middle_name":"",
  "last_name": "",    
  "gender":"",        // male or female
  "date_of_birth":  
  "current_ctc": {{  
    "amount": "",     // The numerical amount 
    "currency": ""    // The currency code 
 }},
  "contact_number": "",        // Phone number with country code if available
  "email": "",
  "occupation": "",            // Occupation of the candidate
  "designation": "",           // latest designation of the candidate
  "full_address": {{
    "address": "",             // Full address of residence
    "country":"",              // Country of residence
    "location": "",            // area of residence
    "city": "",                // City of residence
    "state": ""                // State/region of residence
  }},
  "skills": [],                // List of technical skills mentioned
 
  "education": [
    {{
      "degree": "",            // Degree/qualification (e.g. BSc, MSc, etc.)
      "institution": "",       // Name of university/college/school
      "marks": "",             // CGPA, percentage, or grade
      "education_level": "",   // Level (e.g. Secondary, Undergraduate, Postgraduate)
      "examination_board": "", // Exam board (if given, e.g. CBSE, State Board)
      "medium_of_study": "",   // Medium of study (English, Hindi, etc.)
      "key_skills": [],        // Any key skills highlighted in this course
      "course_type": "",      
       // Allowed values: "Full-time", "Part-time", "Distance".
       // If the resume text does NOT explicitly mention one of these, return exactly "" (empty string).
       // Do not guess, infer, or assume.
      "specialization": "",    // Specialization/major (if mentioned)
      "starting_year": "",     // Year course started
      "ending_year": ""        // Year of completion
    }}
  ],
 
  "projects": [
    {{
      "name": "",                       // Project title
      "project_duration_from": "",      // Start date (YYYY or YYYY-MM)
      "project_duration_to": "",        // End date (YYYY or YYYY-MM)
      "discribe_about_project": "",     // Description as verbatim
      "key_skills": []               // Extract all the Skills applied in project
    }}
  ],
 
  "certification": [
    {{
      "certification_name": "", // Certification name
      "issuer":""   ,            // Issuing authority/platform
      "certification_url": "",  // Certificate link if available
      "valid_from": "",         // Start/issue date
      "valid_till": ""          // Expiry date (if given)
    }}
  ],
 
  "employment": [
    {{
      "company_name": "",                 // Company name
      "designation": "",               // Job title/role id if available
      "start_date": "",                   // Job start date
      "end_date": "",                     // Job end date
      "current_employment": "",           // Yes/No or True/False
      "available_to_work": "",         // Availability status
      "annual_salary": "",                // Annual salary if mentioned
      "describe_what_you_did_at_work": "" // Exact description of tasks as verbatim
    }}
  ],
 
  "internship": [
    {{
      "company_name": "",                 // Internship company name
      "internship_duration_from": "",     // Internship start date
      "internship_duration_to": "",       // Internship end date
      "discribe_about_internship": "",    // Description of internship tasks as verbatim
      "key_skills": [],                   // Skills applied in internship
      "project_url": ""                   // Link to project if any
    }}
  ],
 
   "achievement": [
    {{
      "achievment_type": "",              // e.g. Academic, Professional, Sports, etc.
      "describe_achievment": ""           // Description verbatim from resume
    }}
  ],
 
  "competitive_exams": [
    {{
      "exam_name": "",                    // Name of competitive exam (e.g. GATE, GRE, IELTS)
      "score": ""                         // Score/Rank as mentioned
    }}
  ],
 
  "languages:": {{
    "language": []     // no extra fields , just the name of the language        
  }}
}}
"""
 
    resp = client.chat.complete(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
    )
    return resp.choices[0].message.content.strip()
 
def parse_date(date_str):
    if not date_str or not isinstance(date_str, str):
        return None
    s = date_str.strip().lower()
    if s in {'current','present','now','ongoing','till date','to date','today'}:
        return datetime.datetime.today()
    s = s.replace("–","-").replace("—","-").replace("/", "-")
    for fmt in ("%B %Y", "%b %Y", "%Y-%m", "%m-%Y", "%Y"):
        try:
            dt = datetime.datetime.strptime(s.title() if "%B" in fmt or "%b" in fmt else s, fmt)
            if fmt == "%Y":
                dt = dt.replace(month=1, day=1)
            else:
                dt = dt.replace(day=1)
            return dt
        except ValueError:
            continue
    return None

def merge_date_ranges(ranges):
    intervals = []
    for r in ranges:
        start = parse_date(r.get('start_date'))
        end = parse_date(r.get('end_date'))
        if start and end and end >= start:
            intervals.append((start, end))
    if not intervals:
        return []
    intervals.sort(key=lambda x: x[0])
    merged = [intervals[0]]
    for current_start, current_end in intervals[1:]:
        last_start, last_end = merged[-1]
        if current_start <= last_end:
            merged[-1] = (last_start, max(last_end, current_end))
        else:
            merged.append((current_start, current_end))
    return merged

def calculate_total_experience_months(work_date_ranges):
    merged_ranges = merge_date_ranges(work_date_ranges)
    total_months = 0
    for start, end in merged_ranges:
        diff = (end.year - start.year) * 12 + (end.month - start.month) + 1
        total_months += diff
    return total_months, merged_ranges

def build_date_ranges_from_employment(parsed_info):
    ranges = []
    for job in (parsed_info.get("EMPLOYMENT") or []):
        ranges.append({
            "start_date": (job.get("start_date") or "").strip(),
            "end_date": (job.get("end_date") or "").strip()
        })
    return [r for r in ranges if r.get("start_date") and r.get("end_date")]
 
def build_date_ranges_from_internship(parsed_info):
    ranges = []
    for it in (parsed_info.get("internship") or []):
        ranges.append({
            "start_date": (it.get("internship_duration_from") or "").strip(),
            "end_date": (it.get("internship_duration_to") or "").strip()
        })
    return [r for r in ranges if r.get("start_date") and r.get("end_date")]
 
def serialize_merged_ranges(merged_ranges):
    return [{"start": s.strftime("%Y-%m"), "end": e.strftime("%Y-%m")} for s, e in merged_ranges]
 
def format_experience(total_months):
    years = total_months // 12
    months = total_months % 12
    return {
        "total_experience": f"{years} year{'s' if years != 1 else ''} {months} month{'s' if months != 1 else ''}"
    }

#FUNCTION FOR COMPARE AND UPDATE
def _as_int(name, val):
    try:
        return int(val)
    except (TypeError, ValueError):
        abort(400, description=f"{name} must be an integer")

def generate_pointwise_assessment(resume_text, job_description):
    prompt = f"""
You are an AI assistant comparing a resume with a job description.

Analyze the resume and job description carefully. Identify all mismatches, gaps, or suggestions needed to improve the resume relative to the job description.

Return all output as a JSON array of strings.

Each string must be an action-oriented statement that follows this general format:
"  "<Action> <Category> - <Single Specific Detail>"

Where:
- <Action> is one of: Add, Remove, Change, Update, Clarify, or other suitable verbs.
- <Category> is a label describing what to act on (e.g., Skills, Certificate, Location Preference, Experience, Project, Summary, Education, Contact Information, etc.)
- <Details> clearly explains what to add, remove, change, or update.

Do NOT combine multiple items in one line. Break them into individual suggestions.
"Add Skills - Python, SQL"
"Add Skills - Python", "Add Skills - SQL"
Do NOT recommend removing or deleting any content from the resume — only suggest additions or updates.

Rules:
- Always use this action-oriented format for every suggestion.
- Do not return plain mismatches or bullet points without actions.
- Return ONLY the JSON array of such strings. No extra text or explanation.

Resume:
"{resume_text}"

Job Description:
"{job_description}"
"""
    return call_mistral(prompt)


def extract_text_resume(file_stream, filename):
    filename = filename.lower()
    content = file_stream.read()
    try:
        if filename.endswith('.pdf'):
            reader = PdfReader(BytesIO(content))
            return "\n".join([page.extract_text() or "" for page in reader.pages])
        elif filename.endswith('.docx'):
            doc = Document(BytesIO(content))
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            abort(400, description="Unsupported file format. Use PDF or DOCX.")
    except Exception as e:
        abort(400, description=f"Error reading resume: {str(e)}")

'''def call_mistral(prompt, expect_json=True):
    print("Prompt sent to Mistral:", prompt)

    if not MISTRAL_API_KEY:
        abort(503, description="Model not configured. Missing API key.")
    try:
        messages = [{"role": "user", "content": prompt}]

        
        response = client.chat.complete(
            model=MODEL,
            messages=messages,
            temperature=0.2,
            max_tokens=2000
        )
        print("Mistral raw response:", response)
        content = response.choices[0].message.content.strip()

        # Remove markdown
        content = re.sub(r"^```json\s*", "", content, flags=re.IGNORECASE)
        content = re.sub(r"```$", "", content)

        # 🔥 Fix invalid newlines inside strings
        content = content.replace('\r', '\\r').replace('\n', '\\n')
        return json.loads(content) if expect_json else content
    except Exception as e:
        abort(500, description=f"Mistral API error: {str(e)}")'''

def fix_json_string(content):
    result = []
    in_string = False
    escape = False

    for char in content:
        # Detect start/end of JSON string
        if char == '"' and not escape:
            in_string = not in_string

        # If inside string → escape newline chars only
        if in_string:
            if char == '\n':
                result.append('\\n')
                continue
            elif char == '\r':
                result.append('\\r')
                continue

        # Handle escape characters
        if char == '\\' and not escape:
            escape = True
        else:
            escape = False

        result.append(char)

    return ''.join(result)


def call_mistral(prompt, expect_json=True):
    print("Prompt sent to Mistral:", prompt)

    if not MISTRAL_API_KEY:
        abort(503, description="Model not configured. Missing API key.")

    try:
        response = client.chat.complete(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=2000
        )

        print("Mistral raw response:", response)

        content = response.choices[0].message.content.strip()

        # ✅ Remove markdown wrappers
        content = re.sub(r"^```json\s*", "", content, flags=re.IGNORECASE)
        content = re.sub(r"^```", "", content)
        content = re.sub(r"```$", "", content)

        # ✅ Fix invalid JSON (only inside strings)
        content = fix_json_string(content)

        # Optional debug
        print("Cleaned content:", content)

        return json.loads(content) if expect_json else content

    except Exception as e:
        abort(500, description=f"Mistral API error: {str(e)}")

def flatten_education(education):
    if isinstance(education, dict):
        education = [education]
    if isinstance(education, list):
        lines = []
        for item in education:
            if not isinstance(item, dict):
                text = str(item).replace("_", " ").strip()
                if text:
                    lines.append(text)
                continue

            inst = item.get("INSTITUTION") or item.get("INSTITUTE") or item.get("SCHOOL")
            degree = item.get("DEGREE")
            duration = item.get("DURATION")
            loc = item.get("LOCATION")
            gpa = item.get("GPA")
            coursework = item.get("COURSEWORK")

            head = " — ".join(p for p in [inst, degree, duration] if p)
            tail_bits = [b for b in [loc, gpa, coursework] if b]
            line = head if head else ""
            if tail_bits:
                line = f"{line} — {', '.join(tail_bits)}" if line else ", ".join(tail_bits)
            if line.strip():
                lines.append(line.strip())
        return "\n".join(lines).strip()
    return str(education).replace("_", " ").strip()

def _comma_string(value):
    if isinstance(value, list):
        return ", ".join(str(v).strip() for v in value if v)
    if isinstance(value, str) and value.strip().startswith("["):
        try:
            parsed = json.loads(value.replace("'", '"'))
            if isinstance(parsed, list):
                return ", ".join(str(v).strip() for v in parsed if v)
        except:
            pass
    return (value or "").strip()

def flatten_experience(experience_list):
    if not isinstance(experience_list, list):
        return str(experience_list)
    output = ""
    for job in experience_list:
        role = job.get("ROLE", "")
        company = job.get("COMPANY", "")
        duration = job.get("DURATION", "")
        location = job.get("LOCATION", "")
        description = job.get("DESCRIPTION", "")
        output += f"{role} at {company} ({duration}, {location})\n"
        for line in description.split(". "):
            if line.strip():
                output += f"- {line.strip().rstrip('.')}\n"
        output += "\n"
    return output.strip()

def flatten_projects(projects):
    if not isinstance(projects, list):
        return str(projects)
    output = ""
    for proj in projects:
        title = proj.get("TITLE", "")
        tech = proj.get("TECHNOLOGIES", "")
        about = proj.get("ABOUT", "")
        output += f"{title} ({tech})\n- {about.strip()}\n\n"
    return output.strip()

def remove_markdown_asterisks(text):
    if not isinstance(text, str):
        return text
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'\*', '', text)
    return text.strip()

def clean_fields(fields):
    def clean_text(text):
        return remove_markdown_asterisks(text.strip()) if isinstance(text, str) else text
    cleaned = {}
    for key, value in fields.items():
        if isinstance(value, str):
            cleaned[key] = clean_text(value)
        elif isinstance(value, list):
            cleaned[key] = [clean_text(v) if isinstance(v, str) else v for v in value]
        elif isinstance(value, dict):
            cleaned[key] = {k: clean_text(v) for k, v in value.items()}
        else:
            cleaned[key] = value
    return cleaned

def fill_template_with_fields(fields_dict):
    doc = Document(TEMPLATE_PATH)
    def is_emptyish(val) -> bool:
        if val is None:
            return True
        s = str(val).strip()
        return s == "" or s.lower() in {"none", "n/a", "na", "null", "[]", "{}"}

    header_map = {
        "SUMMARY": "PROFILE",
        "EDUCATION_SECTION": "EDUCATION",
        "TECHNICAL_SKILLS": "TECHNICAL SKILLS",
        "CERTIFICATIONS": "CERTIFICATIONS",
        "EXPERIENCE": "WORK EXPERIENCE",
        "PROJECTS": "PROJECTS",
        "LANGUAGES": "LANGUAGES",
        "PHONE": "CONTACT",
        "EMAIL": "CONTACT",
        
    }

    headers_to_keep = {}
    for field, header in header_map.items():
        headers_to_keep.setdefault(header, False)
        if not is_emptyish(fields_dict.get(field, "")):
            headers_to_keep[header] = True

    def insert_with_format(paragraph, text, *, bold_all=False, uppercase=False, font_size=None):
        if text is None:
            return
        text = str(text)
        if bold_all:
            run = paragraph.add_run(text.upper() if uppercase else text)
            run.bold = True
            if font_size:
                run.font.size = Pt(font_size)
            return
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if not part:
                continue
            clean = part.replace("**", "")
            run = paragraph.add_run(clean.upper() if uppercase else clean)
            if part.startswith("**") and part.endswith("**"):
                run.bold = True
            if font_size:
                run.font.size = Pt(font_size)

    def clear_paragraph(paragraph):
        for r in paragraph.runs:
            r.text = ""

    for para in list(doc.paragraphs):
        text = para.text
        if not text:
            continue
        for key, value in fields_dict.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in text:
                val_str = "" if value is None else str(value).strip()
                if is_emptyish(val_str):
                    p = para._element
                    p.getparent().remove(p)
                    break 

                clear_paragraph(para)
                if key == "FULL_NAME":
                    insert_with_format(para, val_str, bold_all=True, uppercase=True, font_size=16)
                else:
                    insert_with_format(para, val_str)

    headers_to_remove = {h for h, keep in headers_to_keep.items() if not keep}

    def remove_paragraph_at_index(idx):
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)

    i = 0
    while i < len(doc.paragraphs):
        txt = doc.paragraphs[i].text.strip().upper()
        if txt in headers_to_remove:
            remove_paragraph_at_index(i)
            while i < len(doc.paragraphs) and doc.paragraphs[i].text.strip() == "":
                remove_paragraph_at_index(i)
            continue
        i += 1

    filename = f"templated_resume_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(RESUME_FOLDER, filename)
    doc.save(filepath)
    return filename

def normalize_experience(block: str) -> str:
    if not isinstance(block, str):
        return ""
    s = block.strip()
    s = re.sub(r'([A-Z][A-Za-z0-9&.,\s\-]+?\([\w\s–-]+\))',
               r'\n**\1**', s)
    s = re.sub(r'\)\s*-\s*', ')\n', s)
    s = re.sub(r'\.\s+', '.\n- ', s)
    s = re.sub(r'\n+', '\n', s).strip()
    lines = [line.strip() for line in s.split("\n") if line.strip()]
    return "\n\n".join(lines)

def normalize_projects(block: str) -> str:
    if not isinstance(block, str):
        return ""
    s = block.strip()
    s = re.sub(r'([A-Za-z0-9\s\-]+?\([\w\s–-]+\))',
               r'\n**\1**', s)
    s = re.sub(r'\)\s*-\s*Technologies:', r')\nTechnologies:', s)
    s = re.sub(r'\.\s+', '.\n- ', s)
    s = re.sub(r'\n+', '\n', s).strip()
    lines = [line.strip() for line in s.split("\n") if line.strip()]
    return "\n\n".join(lines)

def normalize_education(block: str) -> str:
    if not isinstance(block, str):
        return ""
    s = block.strip()
    s = re.sub(r'([A-Z][A-Za-z0-9&.,\s\-]+?\(\d{4}[-–]?\d{4}\))',
               r'\n**\1**', s)
    s = re.sub(r'([A-Z][A-Za-z0-9&.,\s\-]+)\s*-\s*',
               r'\n**\1**\n', s)
    s = re.sub(r'\)\s*-\s*', ')\n', s)
    s = re.sub(r'\.\s+', '.\n- ', s)
    s = re.sub(r'\n+', '\n', s).strip()
    lines = [line.strip() for line in s.split("\n") if line.strip()]
    return "\n\n".join(lines)

def convert_docx_to_pdf_libreoffice(docx_path: str, out_dir: str = None, timeout: int = 120) -> str:
    docx_path = os.path.abspath(docx_path)
    if out_dir is None:
        out_dir = os.path.dirname(docx_path)
    out_dir = os.path.abspath(out_dir)
    os.makedirs(out_dir, exist_ok=True)

    soffice = shutil.which("soffice")
    if not soffice:
        raise RuntimeError("LibreOffice 'soffice' not found in PATH. Please install LibreOffice.")

    expected_pdf = os.path.join(out_dir, pathlib.Path(docx_path).with_suffix(".pdf").name)
    cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path]
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
    if not os.path.exists(expected_pdf) or os.path.getsize(expected_pdf) == 0:
        raise RuntimeError(
            f"LibreOffice failed to convert DOCX to PDF. "
            f"Stdout: {result.stdout.decode(errors='ignore')}\n"
            f"Stderr: {result.stderr.decode(errors='ignore')}"
        )
    return expected_pdf

def build_file_format(path: str) -> dict:
    st = os.stat(path)
    mime, _ = mimetypes.guess_type(path)
    mime = mime or "application/octet-stream"

    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")

    last_modified = int(st.st_mtime)
    last_modified_iso = datetime.datetime.fromtimestamp(last_modified, tz=datetime.timezone.utc).isoformat()

    return {
        "name": os.path.basename(path),
        "size": st.st_size,
        "type": mime,
        "lastModified": last_modified,
        "lastModifiedDate": last_modified_iso,
        "file": f"data:{mime};base64,{b64}",
    }

# API routes
@app.route("/extract", methods=["POST"])
def extract_resume():
    start_time = time.time()
 
    def log_and_respond(payload, *, http_status=200, status_code=1,
                        message="", ai_response=None, user_input_obj=None, error_message="", metadata_obj=None):
        response_time_ms = int((time.time() - start_time) * 1000)
        try:
            user_input_str = json.dumps(user_input_obj or {}, ensure_ascii=False)
        except Exception:
            user_input_str = "{}"
        try:
            ai_response_str = json.dumps(ai_response if ai_response is not None else {}, ensure_ascii=False)
        except Exception:
            ai_response_str = "{}"
        try:
            metadata_str = json.dumps(metadata_obj or {}, ensure_ascii=False)
        except Exception:
            metadata_str = "{}"
 
        try:
            conn = mysql.connector.connect(**db_config)
            insert_ai_log(
                conn=conn,
                job_id=None,
                user_id=None,
                activity_type="1",
                service_name="mistral",
                model_name=MODEL,
                user_input=user_input_str,
                ai_response_text=ai_response_str,
                response_time_ms=response_time_ms,
                status=status_code,
                error_message=error_message or message or "",
                metadata=metadata_str
            )
            conn.close()
        except Exception as log_err:
            payload = dict(payload)
            payload["_log_warning"] = f"Logging failed: {log_err}"
 
        return jsonify(payload), http_status
 
    uploaded_resume = request.files.get("resume")
    video1 = request.files.get("video1")
    video2 = request.files.get("video2")
 
    if not uploaded_resume:
        return log_and_respond(
            {"error": "Resume file is required"},
            http_status=400,
            status_code=2,
            message="No resume file uploaded",
            user_input_obj={"has_resume": False}
        )
 
    filename = uploaded_resume.filename
    os.makedirs("generated_resumes", exist_ok=True)
    file_path = os.path.join("generated_resumes", filename)
    uploaded_resume.save(file_path)
 
    try:
        resume_text = extract_text(file_path)
    except Exception as e:
        return log_and_respond(
            {"error": str(e)},
            http_status=400,
            status_code=2,
            message="Error reading resume",
            error_message=str(e),
            user_input_obj={"filename": filename}
        )
 
    video_texts = []
    for idx, video in enumerate([video1, video2], start=1):
        if video and video.filename:
            try:
                v_path = save_upload(video, f"video{idx}")
                v_json = transcribe_local_file(v_path)  # AssemblyAI transcription
                video_texts.append(v_json.get("text", ""))
                os.remove(v_path)
            except Exception as ve:
                video_texts.append(f"[Video{idx} transcription failed: {ve}]")
        else:
            video_texts.append("")
 
    video_transcripts_payload = {
        "video1": video_texts[0] if len(video_texts) > 0 else "",
        "video2": video_texts[1] if len(video_texts) > 1 else ""
    }
 
    combined_video_text = "\n".join(v for v in video_texts if v)
 
    try:
        verdict = is_resume_with_mistral(resume_text)
        if verdict != "yes":
            return log_and_respond(
                {"message": "Upload a valid resume."},
                http_status=400,
                status_code=2,
                message="Upload a valid resume.",
                user_input_obj={"filename": filename, "verdict": verdict},
                ai_response={"verdict": verdict}
            )
 
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            embedded_links = extract_links_from_pdf(file_path)
        elif ext == ".docx":
            embedded_links = extract_links_from_docx(file_path)
        else:
            embedded_links = []
 
        visible_links = extract_urls_from_text(resume_text)
        all_links = list(set(embedded_links + visible_links))
        social_links = extract_social_links(all_links)
 
        raw_info = extract_info_with_mistral(resume_text, combined_video_text)
        parsed_info = clean_response_json(raw_info)
        parsed_info = parsed_info if isinstance(parsed_info, dict) else {}
 
        for edu in parsed_info.get("education", []):
            if edu.get("starting_year") and edu.get("starting_year") == edu.get("ending_year"):
                edu["starting_year"] = ""
 
        emp_ranges_raw = build_date_ranges_from_employment(parsed_info)
        int_ranges_raw = build_date_ranges_from_internship(parsed_info)
 
        emp_months, emp_merged = calculate_total_experience_months(emp_ranges_raw)
        int_months, int_merged = calculate_total_experience_months(int_ranges_raw)
 
        total_months = emp_months + int_months
 
        experience_payload = {
            "employment": {**format_experience(emp_months)},
            "internship": {**format_experience(int_months)},
            "total": format_experience(total_months)
        }
 
        payload = {
            "status": 1,
            "data": {
                "info": parsed_info,
                "links": social_links,
                "experience": experience_payload,
                "video_transcripts": video_transcripts_payload
            }
        }
 
        return log_and_respond(
            payload,
            http_status=200,
            status_code=1,
            user_input_obj={"filename": filename, "verdict": "yes"},
            ai_response=payload,
            metadata_obj={
                "embedded_links": embedded_links,
                "visible_links": visible_links,
                "employment_merged_ranges": serialize_merged_ranges(emp_merged),
                "internship_merged_ranges": serialize_merged_ranges(int_merged)
            }
        )
 
    except Exception as e:
        return log_and_respond(
            {"error": str(e)},
            http_status=500,
            status_code=2,
            message="Internal error",
            error_message=str(e),
            user_input_obj={"filename": filename}
        )
 
@app.route("/compare-resume", methods=["POST"])
def compare_resume():
    start_time = time.time()
    db_conn = None
    cursor = None

    def log_and_respond(msg, *, http_status=400, job_id=None, user_id=None, err_status=2):
        response_time_ms = int((time.time() - start_time) * 1000)
        try:
            db_conn_log = mysql.connector.connect(**db_config)
            insert_ai_log(
                conn=db_conn_log,
                job_id=job_id,
                user_id=user_id,
                activity_type="5",
                service_name="mistral",
                model_name=MODEL,
                user_input=json.dumps({"job_id": job_id, "user_id": user_id}),
                ai_response_text="",
                response_time_ms=response_time_ms,
                status=err_status,
                error_message=msg,
                metadata=json.dumps({})
            )
            db_conn_log.close()
        except Exception as log_err:
            print("Logging failed:", log_err)

        return jsonify({"status": 0, "error": {"message": msg}}), http_status

    try:
        payload = request.get_json(silent=True) or {}
        job_id = (
            request.form.get("job_id")
            or payload.get("job_id")
            or request.args.get("job_id")
        )
        user_id = (
            request.form.get("user_id")
            or payload.get("user_id")
            or request.args.get("user_id")
        )

        job_id = _as_int("job_id", job_id)
        user_id = _as_int("user_id", user_id)

        if not job_id or not user_id:
            return log_and_respond("Both job_id and user_id are required",
                                   http_status=400, job_id=job_id, user_id=user_id)

        if not MISTRAL_API_KEY:
            return log_and_respond("Model not configured. Missing API key.",
                                   http_status=503, job_id=job_id, user_id=user_id)

        db_conn = mysql.connector.connect(**db_config)
        db_conn.start_transaction()
        cursor = db_conn.cursor()

        cursor.execute("""
            SELECT COALESCE(cost_per_response, 0)
            FROM tbl_master_package
            WHERE id = 2
            LIMIT 1
        """)
        row = cursor.fetchone()
        if not row:
            return log_and_respond("Pricing not configured. Please contact support.",
                                   http_status=500, job_id=job_id, user_id=user_id)

        cost_per_response = Decimal(str(row[0] or "0"))

        cursor.execute("""
            SELECT COALESCE(total_credits, 0)
            FROM tbl_user_account
            WHERE user_id = %s
            FOR UPDATE
        """, (user_id,))
        row = cursor.fetchone()
        if not row:
            return log_and_respond("Account not found for user.",
                                   http_status=404, job_id=job_id, user_id=user_id)

        total_credits = Decimal(str(row[0] or "0"))
        if total_credits < cost_per_response:
            db_conn.rollback()
            return jsonify({
                "status": 0,
                "error": {
                    "message": "Insufficient credits",
                    "required": str(cost_per_response),
                    "available": str(total_credits)
                }
            }), 402

        cursor.execute("SELECT description_plain_text FROM tbl_job WHERE id = %s", (job_id,))
        job_result = cursor.fetchone()
        if not job_result:
            return log_and_respond("Job description not found for given job_id",
                                   http_status=404, job_id=job_id, user_id=user_id)

        job_description = (job_result[0] or "").strip()
        if len(job_description) < 30:
            return log_and_respond("Job description is too short.",
                                   http_status=422, job_id=job_id, user_id=user_id)

        cursor.execute("SELECT resume_file_name FROM tbl_job_seeker_resume WHERE user_id = %s", (user_id,))
        resume_result = cursor.fetchone()
        if not resume_result:
            return log_and_respond("Resume filename not found for given user_id",
                                   http_status=404, job_id=job_id, user_id=user_id)

        resume_filename = resume_result[0]
        resume_path = os.path.join(RESUME_FOLDER, resume_filename)
        if not os.path.exists(resume_path):
            return log_and_respond(f"Resume file '{resume_filename}' not found in resumes folder",
                                   http_status=404, job_id=job_id, user_id=user_id)

        with open(resume_path, "rb") as f:
            resume_text = extract_text_resume(f, resume_filename) or ""

        if len(resume_text.strip()) < 50:
            return log_and_respond("Resume text is too short or unreadable.",
                                   http_status=422, job_id=job_id, user_id=user_id)

        assessment = generate_pointwise_assessment(resume_text, job_description)
        if not isinstance(assessment, list):
            assessment = [str(assessment)]
        assessment = [x.strip() for x in assessment if str(x).strip()]

        response_time_ms = int((time.time() - start_time) * 1000)
        insert_ai_log(
            conn=db_conn,
            job_id=job_id,
            user_id=user_id,
            activity_type="4",
            service_name="mistral",
            model_name=MODEL,
            user_input=json.dumps({
                "job_id": job_id,
                "user_id": user_id,
                "resume_text": resume_text[:1000],
                "job_description": job_description[:1000]
            }),
            ai_response_text=json.dumps(assessment),
            response_time_ms=response_time_ms,
            status=1,
            error_message="",
            metadata=json.dumps({})
        )

        cursor.execute(
            "UPDATE tbl_user_account SET total_credits = total_credits - %s WHERE user_id = %s",
            (cost_per_response, user_id)
        )

        cursor.execute("SELECT COALESCE(total_credits, 0) FROM tbl_user_account WHERE user_id = %s", (user_id,))
        remaining = Decimal(str(cursor.fetchone()[0] or "0"))

        timestamp_now = int(time.time())
        cursor.execute("""
            INSERT INTO tbl_user_credit_debit_usage
                (user_id, user_type, credits, usage_type, used_on, balance_credits, type, user_credit_purchase_id)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            user_id, 2, cost_per_response, 2,
            timestamp_now, remaining, 2, None
        ))
        usage_id = cursor.lastrowid

        cursor.execute("""
            INSERT INTO tbl_user_credit_debit_usage_activity
                (user_id, credits_used, balance_credits, user_credit_usage_id, job_id)
            VALUES (%s, %s, %s, %s, %s)
        """, (
            user_id, cost_per_response, remaining, usage_id, job_id
        ))

        db_conn.commit()

        cursor.close()
        db_conn.close()

        return jsonify({
            "status": 1,
            "data": assessment,
            "message": "Recommendation has been received successfully"
        }), 200

    except Exception as e:
        traceback.print_exc()
        try:
            if db_conn:
                db_conn.rollback()
        except:
            pass

        try:
            response_time_ms = int((time.time() - start_time) * 1000)
            db_conn_log = mysql.connector.connect(**db_config)
            insert_ai_log(
                conn=db_conn_log,
                job_id=locals().get("job_id"),
                user_id=locals().get("user_id"),
                activity_type="4",
                service_name="mistral",
                model_name=MODEL,
                user_input=json.dumps({
                    "job_id": locals().get("job_id", "missing"),
                    "user_id": locals().get("user_id", "missing")
                }),
                ai_response_text="",
                response_time_ms=response_time_ms,
                status=2,
                error_message=str(e),
                metadata=json.dumps({})
            )
            db_conn_log.close()
        except Exception as log_err:
            print("Logging failed in exception handler:", log_err)

        if cursor:
            cursor.close()
        if db_conn:
            db_conn.close()

        return jsonify({
            "status": 0,
            "error": {"message": "Internal server error. Please try again later."}
        }), 500

@app.route("/update-resume", methods=["POST"])
def update_resume():
    def log_and_respond(message, *, http_status=400, job_id=None, user_id=None, data=None, err_status=2):
        response_time_ms = int((time.time() - start_time) * 1000)
        try:
            db_conn_log = mysql.connector.connect(**db_config)
            insert_ai_log(
                conn=db_conn_log,
                job_id=job_id,
                user_id=user_id,
                activity_type="5",
                service_name="mistral",
                model_name=MODEL,
                user_input=json.dumps({
                    "job_id": job_id,
                    "user_id": user_id,
                    "recommendations": (data or {}).get("recommendations", [])
                }, ensure_ascii=False),
                ai_response_text="",
                response_time_ms=response_time_ms,
                status=err_status,
                error_message=message,
                metadata=json.dumps({}, ensure_ascii=False)
            )
            try:
                db_conn_log.close()
            except:
                pass
        except Exception as log_err:
            print("Logging failed:", log_err)

        return jsonify({"status": 0, "error": {"message": message}}), http_status

    start_time = time.time()
    db_conn = None
    cursor = None

    try:
        if not request.is_json:
            return log_and_respond("Body must be JSON.", http_status=400)

        data = request.json or {}
        user_id_in = data.get("user_id")
        job_id_in = data.get("job_id")
        recommendations = data.get("recommendations", [])

        if user_id_in in (None, "", []) or job_id_in in (None, "", []):
            return log_and_respond("Both user_id and job_id are required.", http_status=400, data=data)

        try:
            user_id = _as_int("user_id", user_id_in)
            job_id = _as_int("job_id", job_id_in)
        except Exception:
            return log_and_respond("user_id and job_id must be integers.", http_status=400, data=data)

        if not isinstance(recommendations, list) or len(recommendations) == 0:
            return log_and_respond("'recommendations' must be a non-empty list.", http_status=400, data=data)

        if not MISTRAL_API_KEY:
            return log_and_respond("Model not configured. Missing API key.", http_status=503, job_id=job_id, user_id=user_id, data=data)

        db_conn = mysql.connector.connect(**db_config)
        db_conn.start_transaction()
        cursor = db_conn.cursor()

        cursor.execute("""
            SELECT COALESCE(cost_per_response, 0)
            FROM tbl_master_package
            WHERE id = 3
            LIMIT 1
        """)
        row = cursor.fetchone()
        if not row:
            db_conn.rollback()
            cursor.close(); db_conn.close()
            cursor = None; db_conn = None
            return log_and_respond("Pricing not configured. Please contact support.", http_status=500, job_id=job_id, user_id=user_id, data=data)

        cost_update = row[0]
        if not isinstance(cost_update, Decimal):
            cost_update = Decimal(str(cost_update or "0"))

        cursor.execute("""
            SELECT COALESCE(total_credits, 0)
            FROM tbl_user_account
            WHERE user_id = %s
            FOR UPDATE
        """, (user_id,))
        row = cursor.fetchone()
        if not row:
            db_conn.rollback()
            cursor.close(); db_conn.close()
            cursor = None; db_conn = None
            return log_and_respond("Account not found for user.", http_status=404, job_id=job_id, user_id=user_id, data=data)

        total_credits = row[0]
        if not isinstance(total_credits, Decimal):
            total_credits = Decimal(str(total_credits or "0"))

        if total_credits < cost_update:
            db_conn.rollback()
            cursor.close(); db_conn.close()
            cursor = None; db_conn = None
            return jsonify({"status": 0, "error": {"message": "Insufficient credits"}}), 402

        cursor.execute("SELECT description_plain_text FROM tbl_job WHERE id = %s", (job_id,))
        job_result = cursor.fetchone()
        if not job_result:
            db_conn.rollback()
            cursor.close(); db_conn.close()
            cursor = None; db_conn = None
            return log_and_respond("Job description not found for given job_id", http_status=404, job_id=job_id, user_id=user_id, data=data)
        job_description = job_result[0] or ""

        cursor.execute("SELECT resume_file_name FROM tbl_job_seeker_resume WHERE user_id = %s", (user_id,))
        resume_result = cursor.fetchone()
        if not resume_result:
            db_conn.rollback()
            cursor.close(); db_conn.close()
            cursor = None; db_conn = None
            return log_and_respond("Resume filename not found for given user_id", http_status=404, job_id=job_id, user_id=user_id, data=data)
        resume_filename = resume_result[0]

        if not resume_filename:
            db_conn.rollback()
            cursor.close(); db_conn.close()
            cursor = None; db_conn = None
            return log_and_respond("Resume filename is empty for user.", http_status=422, job_id=job_id, user_id=user_id, data=data)

        resume_path = os.path.join(RESUME_FOLDER, resume_filename)
        if not os.path.exists(resume_path):
            db_conn.rollback()
            cursor.close(); db_conn.close()
            cursor = None; db_conn = None
            return log_and_respond(f"Resume file '{resume_filename}' not found in resumes folder", http_status=404, job_id=job_id, user_id=user_id, data=data)

        with open(resume_path, "rb") as f:
            resume_text = extract_text_resume(f, resume_filename) or ""

        if len(resume_text.strip()) < 50:
            db_conn.rollback()
            cursor.close(); db_conn.close()
            cursor = None; db_conn = None
            return log_and_respond("Resume text is too short or unreadable.", http_status=422, job_id=job_id, user_id=user_id, data=data)

        if len((job_description or "").strip()) < 30:
            db_conn.rollback()
            cursor.close(); db_conn.close()
            cursor = None; db_conn = None
            return log_and_respond("Job description is too short.", http_status=422, job_id=job_id, user_id=user_id, data=data)

        if not os.path.exists(TEMPLATE_PATH):
            db_conn.rollback()
            cursor.close(); db_conn.close()
            cursor = None; db_conn = None
            return log_and_respond("Server template missing. Notify support.", http_status=500, job_id=job_id, user_id=user_id, data=data)

        prompt = f"""
You are an expert professional resume writer.

TASK:
- Take the given Resume text.
- Incorporate ONLY the items explicitly listed in "Recommendations".
- You MAY enhance the clarity, grammar, and phrasing of existing content in the Resume text to make it more professional and polished.
- Do NOT generate new content that is not in the Resume or in Recommendations.
- Do NOT change existing content except to add the Recommendations in a natural, professional way.
- The output must be a FINAL JSON object, strictly matching the schema below.

Output schema (NOTE: return *strings*, not arrays):
{{
  "FULL_NAME": "",
  "SUMMARY": "",
  "EDUCATION": "Single or multi-line human-readable text combining institution, degree, duration, location, GPA, coursework as appropriate.",
  "TECHNICAL_SKILLS": "Comma-separated list (single string). No brackets.",
  "CERTIFICATIONS": "Comma-separated list (single string). No brackets.",
  "EXPERIENCE": "Flattened, polished resume text. Use bullet-style lines with leading '- ' where helpful. No JSON arrays.",
  "PROJECTS": "Flattened, polished resume text. Use bullet-style lines with leading '- ' where helpful. No JSON arrays.",
  "LANGUAGES": "Comma-separated list (single string).",
  "PHONE": "",
  "EMAIL": ""
}}

Guidelines:
- Use the Resume text as the base.
- Add each Recommendation explicitly and realistically into the most appropriate section.
- Do NOT add or invent new skills, projects, or experience unless explicitly in Recommendations.
- Keep formatting consistent with a professional resume.
- Do NOT fabricate new content beyond Resume + Recommendations.

Resume:
{resume_text}

Recommendations (MUST be incorporated):
{json.dumps(recommendations)}
"""
        filled_fields = call_mistral(prompt)
        print("Raw filled fields from Mistral:", filled_fields)

        filled_fields["EXPERIENCE"] = normalize_experience(filled_fields.get("EXPERIENCE", ""))
        filled_fields["PROJECTS"] = normalize_projects(filled_fields.get("PROJECTS", ""))
        filled_fields["EDUCATION"] = normalize_education(filled_fields.get("EDUCATION", ""))

        if "EDUCATION" in filled_fields and "EDUCATION_SECTION" not in filled_fields:
            filled_fields["EDUCATION_SECTION"] = filled_fields.get("EDUCATION", "")

        for key in ["TECHNICAL_SKILLS", "CERTIFICATIONS", "LANGUAGES"]:
            filled_fields[key] = _comma_string(filled_fields.get(key, ""))

        if not isinstance(filled_fields.get("EXPERIENCE", ""), str):
            filled_fields["EXPERIENCE"] = flatten_experience(filled_fields.get("EXPERIENCE", ""))

        if not isinstance(filled_fields.get("PROJECTS", ""), str):
            filled_fields["PROJECTS"] = flatten_projects(filled_fields.get("PROJECTS", ""))

        filled_fields = clean_fields(filled_fields)

        templating_fields = {
            "FULL_NAME": filled_fields.get("FULL_NAME", ""),
            "SUMMARY": filled_fields.get("SUMMARY", ""),
            "EDUCATION_SECTION": filled_fields.get("EDUCATION_SECTION", ""),
            "TECHNICAL_SKILLS": filled_fields.get("TECHNICAL_SKILLS", ""),
            "CERTIFICATIONS": filled_fields.get("CERTIFICATIONS", ""),
            "EXPERIENCE": filled_fields.get("EXPERIENCE", ""),
            "PROJECTS": filled_fields.get("PROJECTS", ""),
            "LANGUAGES": filled_fields.get("LANGUAGES", ""),
            "PHONE": filled_fields.get("PHONE", ""),
            "EMAIL": filled_fields.get("EMAIL", "")
        }

        final_docx = fill_template_with_fields(templating_fields)
        docx_path = os.path.join(RESUME_FOLDER, final_docx)

        try:
            pdf_path = convert_docx_to_pdf_libreoffice(docx_path, RESUME_FOLDER)
            cv_base_path = pdf_path
            respData = {"curriculum_vitae": os.path.basename(cv_base_path)}
            cv_format = build_file_format(cv_base_path)
            respData["curriculum_vitae_format"] = cv_format
        except Exception as e:
            print("❌ PDF conversion error:", str(e))  # 🔥 IMPORTANT

            db_conn.rollback()

            try:
                cursor.close()
            except:
                pass

            try:
                db_conn.close()
            except:
                pass

            cursor = None
            db_conn = None

            return log_and_respond(
                f"PDF conversion failed: {str(e)}",  # 🔥 send actual error
                http_status=500,
                job_id=job_id,
                user_id=user_id,
                data=data
            )
        upcur = db_conn.cursor()
        try:
            upcur.execute("""
                INSERT INTO tbl_job_application_resume
                    (job_id, user_id, resume, resume_current_status, file_data)
                VALUES
                    (
                      %s, %s, %s, %s,
                      JSON_OBJECT(
                        'file', %s,
                        'lastModified', %s,
                        'lastModifiedDate', %s,
                        'name', %s,
                        'size', %s,
                        'type', %s
                      )
                    )
                ON DUPLICATE KEY UPDATE
                    resume = VALUES(resume),
                    resume_current_status = VALUES(resume_current_status),
                    file_data = VALUES(file_data)
            """, (
                job_id, user_id, os.path.basename(pdf_path), 1,
                cv_format.get('file'), cv_format.get('lastModified'), cv_format.get('lastModifiedDate'),
                cv_format.get('name'), cv_format.get('size'), cv_format.get('type'),
            ))
            app_resume_id = upcur.lastrowid

            cur_demote = db_conn.cursor()
            try:
                cur_demote.execute("""
                    UPDATE tbl_job_application_resume
                    SET resume_current_status = 2
                    WHERE user_id = %s
                    AND job_id = %s
                    AND id <> %s
                    AND resume_current_status <> 2
                """, (user_id, job_id, app_resume_id))
            finally:
                cur_demote.close()

            db_conn.commit()
        finally:
            upcur.close()

        rec_cur = db_conn.cursor()
        try:
            rec_cur.execute("""
                SELECT recommendation
                FROM tbl_job_application_resume_recommendation
                WHERE job_application_resume_id = %s
            """, (app_resume_id,))
            existing = {row[0] for row in rec_cur.fetchall()}
            rows = [
                (app_resume_id, rec.strip())
                for rec in (recommendations or [])
                if isinstance(rec, str) and rec.strip() and rec.strip() not in existing
            ]
            if rows:
                rec_cur.executemany("""
                    INSERT INTO tbl_job_application_resume_recommendation
                        (job_application_resume_id, recommendation)
                    VALUES (%s, %s)
                """, rows)
                db_conn.commit()
        finally:
            rec_cur.close()

        cursor = db_conn.cursor()
        cursor.execute(
            "UPDATE tbl_user_account SET total_credits = total_credits - %s WHERE user_id = %s",
            (cost_update, user_id)
        )
        db_conn.commit()

        cursor.execute("SELECT total_credits FROM tbl_user_account WHERE user_id = %s", (user_id,))
        row_balance = cursor.fetchone()
        balance_after_update = Decimal(str(row_balance[0])) if row_balance else Decimal("0")

        cursor.close(); db_conn.close()
        cursor = None; db_conn = None

        scoring_done = False
        score_result = None
        cost_scoring = Decimal("0")
        balance_after_scoring = balance_after_update

        db_conn2 = mysql.connector.connect(**db_config)
        try:
            db_conn2.start_transaction()
            cur2 = db_conn2.cursor()
            cur2.execute("""
                SELECT COALESCE(cost_per_response, 0)
                FROM tbl_master_package
                WHERE id = 3
                LIMIT 1
            """)
            row = cur2.fetchone()
            if row:
                cost_scoring = Decimal(str(row[0] or "0"))

                cur2.execute("""
                    SELECT COALESCE(total_credits, 0)
                    FROM tbl_user_account
                    WHERE user_id = %s
                    FOR UPDATE
                """, (user_id,))
                row = cur2.fetchone()
                can_afford = (row is not None) and (Decimal(str(row[0])) >= cost_scoring)

                cur2.close(); db_conn2.rollback()

                if can_afford:
                    score_result = run_scoring_for_pair(user_id, job_id)
                    if not (isinstance(score_result, dict) and "error" in score_result):
                        db_conn3 = mysql.connector.connect(**db_config)
                        try:
                            db_conn3.start_transaction()
                            cur3 = db_conn3.cursor()
                            cur3.execute("""
                                SELECT COALESCE(total_credits, 0)
                                FROM tbl_user_account
                                WHERE user_id = %s
                                FOR UPDATE
                            """, (user_id,))
                            row3 = cur3.fetchone()
                            if row3 and Decimal(str(row3[0])) >= cost_scoring:
                                cur3.execute(
                                    "UPDATE tbl_user_account SET total_credits = total_credits - %s WHERE user_id = %s",
                                    (cost_scoring, user_id)
                                )
                                db_conn3.commit()

                                cur3.execute("SELECT total_credits FROM tbl_user_account WHERE user_id = %s", (user_id,))
                                row_bal2 = cur3.fetchone()
                                balance_after_scoring = Decimal(str(row_bal2[0])) if row_bal2 else Decimal("0")

                                scoring_done = True
                            db_conn3.commit()
                            cur3.close(); db_conn3.close()
                        except Exception:
                            try:
                                db_conn3.rollback()
                            except:
                                pass
                            try:
                                cur3.close()
                            except:
                                pass
                            try:
                                db_conn3.close()
                            except:
                                pass
            try:
                db_conn2.rollback()
            except:
                pass
            try:
                cur2.close()
            except:
                pass
            try:
                db_conn2.close()
            except:
                pass
        except Exception:
            try:
                db_conn2.rollback()
            except:
                pass
            try:
                cur2.close()
            except:
                pass
            try:
                db_conn2.close()
            except:
                pass

        try:
            db_conn_usage = mysql.connector.connect(**db_config)
            db_conn_usage.start_transaction()
            cur_usage = db_conn_usage.cursor()

            used_on_ts_update = int(datetime.datetime.now().timestamp())

            cur_usage.execute("""
                INSERT INTO tbl_user_credit_debit_usage
                    (user_id, user_type, credits, usage_type, used_on, balance_credits, type)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """, (user_id, 2, cost_update, 3, used_on_ts_update, balance_after_update, 2))
            db_conn_usage.commit()
            usage_id_update = cur_usage.lastrowid

            cur_usage.execute("""
                INSERT INTO tbl_user_credit_debit_usage_activity
                    (user_credit_usage_id, user_id, job_id, credits_used, balance_credits)
                VALUES (%s, %s, %s, %s, %s)
            """, (usage_id_update, user_id, job_id, cost_update, balance_after_update))
            db_conn_usage.commit()

            if scoring_done:
                used_on_ts_score = int(datetime.datetime.now().timestamp())
                cur_usage.execute("""
                    INSERT INTO tbl_user_credit_debit_usage
                        (user_id, user_type, credits, usage_type, used_on, balance_credits, type)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (user_id, 2, cost_scoring, 6, used_on_ts_score, balance_after_scoring, 2))
                db_conn_usage.commit()
                usage_id_score = cur_usage.lastrowid

                cur_usage.execute("""
                    INSERT INTO tbl_user_credit_debit_usage_activity
                        (user_credit_usage_id, user_id, job_id, credits_used, balance_credits)
                    VALUES (%s, %s, %s, %s, %s)
                """, (usage_id_score, user_id, job_id, cost_scoring, balance_after_scoring))
                db_conn_usage.commit()

            cur_usage.close(); db_conn_usage.close()
        except Exception as usage_err:
            print("⚠️ Credit usage/activity insert failed:", usage_err)
            try:
                db_conn_usage.rollback()
            except:
                pass
            try:
                cur_usage.close()
            except:
                pass
            try:
                db_conn_usage.close()
            except:
                pass

        response_time_ms = int((time.time() - start_time) * 1000)
        respData_for_log = copy.deepcopy(respData)
        try:
            if "curriculum_vitae_format" in respData_for_log:
                respData_for_log["curriculum_vitae_format"].pop("file", None)
        except Exception:
            pass

        def _truncate_bytes(s: str, max_bytes: int) -> str:
            if s is None:
                return ""
            b = s.encode("utf-8", errors="ignore")
            if len(b) <= max_bytes:
                return s
            return b[: max_bytes - 3].decode("utf-8", errors="ignore") + "..."

        user_input_str = json.dumps({
            "job_id": job_id,
            "user_id": user_id,
            "recommendations": recommendations
        }, ensure_ascii=False)

        ai_response_str = json.dumps(respData_for_log, ensure_ascii=False)
        metadata_str = json.dumps({
            "cost_update_id3": str(cost_update),
            "scoring_debited": bool(scoring_done)
        }, ensure_ascii=False)

        user_input_str = _truncate_bytes(user_input_str, 8000)
        ai_response_str = _truncate_bytes(ai_response_str, 60000)
        metadata_str = _truncate_bytes(metadata_str, 8000)

        try:
            db_conn_log = mysql.connector.connect(**db_config)
            insert_ai_log(
                conn=db_conn_log,
                job_id=job_id,
                user_id=user_id,
                activity_type="5",
                service_name="mistral",
                model_name=MODEL,
                user_input=user_input_str,
                ai_response_text=ai_response_str,
                response_time_ms=response_time_ms,
                status=1,
                error_message="",
                metadata=metadata_str
            )
            db_conn_log.close()
        except Exception as log_err:
            print("Success logging failed:", log_err)

        if scoring_done and isinstance(score_result, dict):
            respData["scoring"] = score_result
        elif "scoring" not in respData:
            respData["scoring"] = {"status": 0, "error": "Scoring not performed"}

        return jsonify({
            "status": 1,
            "data": respData,
            "message": "Resume updated successfully"
        }), 200

    except Exception as e:
        traceback.print_exc()
        try:
            if db_conn:
                db_conn.rollback()
        except:
            pass
        try:
            if cursor:
                cursor.close()
        except:
            pass
        try:
            if db_conn:
                db_conn.close()
        except:
            pass
        return jsonify({"status": 0, "error": {"message": "Internal server error. Please try again later."}}), 500


#4 for hr and seeker scores
RESUMES_FILE="generated_resumes"
#used to extract text from the file_path in /init and /process_job
def extract_text_from_file(file_path):
    if file_path.endswith('.pdf'):
        reader = PdfReader(file_path)
        return " ".join(page.extract_text() or "" for page in reader.pages)
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return " ".join(para.text for para in doc.paragraphs) 
    return "error reading file in extract_text_from_file  "

#/init and /process_job called in get_all_mistral_scores() function 
def extract_json_block(text):
    try:
        json_text = re.search(r'\{[\s\S]+\}', text).group()
        return json.loads(json_text)
    except Exception as e:
        #print(f"⚠️ Failed to extract valid JSON from Mistral response: {text}")
        return  "Failed to extract valid JSON from Mistral response"

# overall score finding using mistral /init and /process_job
def get_all_mistral_scores(resume_text, parameters_weightages_dict):
    prompt_parts = []

    for parameter, weightages in parameters_weightages_dict.items():
        scoring_levels = "\n".join([f"{i} -> {weightages.get(i, '')}" for i in range(10, -1, -1)])
        part = f"""
Parameter: {parameter}

Scoring Levels:
{scoring_levels}

"""
        #print(part)mistral_client.chat.completions.create
        prompt_parts.append(part)

    final_prompt = f"""
You are an expert resume evaluator.

You will be provided with:

A candidate's resume in plain text format.

Multiple evaluation parameters.

Each parameter includes numeric scoring levels from 10 to 0.

Strict evaluation rules:

For Years of Experience:
• Extract the number strictly from explicit mentions of work experience in the resume text (e.g., “7 years of experience”, “8+ years”, “over 10 years”).
• Ignore any decimal part (e.g., 7.8 → 7, 8.9 → 8).
• Match the resulting integer EXACTLY against the scoring levels.
• If no valid number of years is found or does not match exactly, assign 0.

For Skills-related fields:
• ONLY check inside the skills section of the resume text (if present).
• Identify all possible matches of keywords/phrases defined in the scoring table.
• Assign ONLY the single highest score corresponding to the strongest match.
• If no skills match, assign 0.

For all other parameters:
• Read and fully scan the entire resume text before scoring.
• Identify all possible matches of keywords/phrases across the entire resume.
• From the matches, assign only the single highest score that corresponds to the strongest match.
• If no keywords/phrases match, assign 0.

NEVER hallucinate or fabricate values. If it’s not explicitly in the resume text, treat it as absent.

Return ONLY a JSON object where keys are parameter names and values are numeric scores (0–10).
Do not add explanations, commentary, or extra text.

Resume:
\"\"\"
{resume_text}
\"\"\"

Evaluate the following:
{''.join(prompt_parts)}

Strictly return this format:
{{ "parameter1": score1, "parameter2": score2, ... }}
"""

    try:
        response = client.chat.complete(
            model=MODEL,
            messages=[{"role": "user", "content": [{"type": "text", "text": final_prompt}]}],
            temperature=0.0,
        )
        content = response.choices[0].message.content.strip()
        scores = extract_json_block(content)
        valid_scores = {k: int(v) for k, v in scores.items() if isinstance(v, int) and 0 <= v <= 10}
        return valid_scores

    except Exception as e:
        #print(f"❌ Mistral API error during bulk evaluation: {e}")
        return {}
    
#get skill ,experience match using the resume text and description /init and /process_job
def get_skill_experience_match(resume_text, job_description):
    prompt = f"""
You are a job matching expert.

Your task is to compare the following Job Description with the Resume Text and return two scores (0-100):

Scoring rules:

1. "skills_score":
   - Do not rely on exact keyword matching.
   - Evaluate how similar the candidate's skills and tools are to those in the job description.
   - If the resume skills are very close to the job requirements, assign a high score.
   - The more the resume skills deviate, the lower the score.

2. "experience_score":
   - Extract the candidate’s years of experience and the required years from the job description.
   - Do not use a fixed formula.
   - If the candidate’s experience is very close to the job requirement, assign a high score.
   - If the candidate has much less or much more than required, reduce the score proportionally.
   - The closer the candidate’s experience to the required, the higher the score.

Return strictly in this JSON format:
{{
  "skills_score": <int>,
  "experience_score": <int>
}}

Job Description:
\"\"\"
{job_description}
\"\"\"

Resume Text:
\"\"\"
{resume_text}
\"\"\"
"""


    try:
        response = client.chat.complete(
            model=MODEL,
            messages=[{"role": "user", "content": [{"type": "text", "text": prompt}]}],
            temperature=0.0,
        )
        content = response.choices[0].message.content.strip()

        # ✅ Extract JSON using regex
        json_match = re.search(r'\{[\s\S]*?\}', content)
        if json_match:
            return json.loads(json_match.group())

        #print(f"⚠️ No valid JSON found in Mistral response: {content}")
        return {"skills_score": 0, "experience_score": 0}

    except Exception as e:
        #print(f"❌ Error in skill/experience match: {e}")
        return {"skills_score": 0, "experience_score": 0}
    


#function to call the ai functions to extract overall_score and skill experience match and store in output table called in both /init and /process_job
def process_seeker(job_id, user_id, resume_text, weightage_rows, job_description):
    process_seeker_conn = None
    process_seeker_cursor= None
    try:
        process_seeker_conn = connection_pool.get_connection()
        process_seeker_cursor = process_seeker_conn.cursor(dictionary=True)
        detailed_scores = {}
        raw_score_total = 0
        total_parameters = len(weightage_rows)

        parameters_weightages_dict = {
            row['parameter']: {i: row.get(f'weightage_{i}') for i in range(11)}
            for row in weightage_rows
        }
        
        scores = get_all_mistral_scores(resume_text, parameters_weightages_dict)

        for parameter, score in scores.items():
            percentage = (score * 10) / total_parameters if total_parameters else 0
            detailed_scores[parameter] = percentage
            raw_score_total += score

        overall_score = sum(detailed_scores.values())
        match_result = get_skill_experience_match(resume_text, job_description)
        # 🔹 Insert new application
        
        process_seeker_cursor.execute("""
                INSERT INTO tbl_job_application (
                    job_id, user_id, overall_score, overall_percent,
                    skill_match, experience_match,
                    created_by, created_at, updated_by, updated_at
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                job_id, user_id, raw_score_total, int(overall_score),
                match_result.get("skills_score", 0),
                match_result.get("experience_score", 0),
                1, int(time.time()), 1, int(time.time())
            ))
        application_id = process_seeker_cursor.lastrowid

        # 🔹 Insert fresh weightages
        
        for parameter, divided_percent in detailed_scores.items():
            process_seeker_cursor.execute(
                "SELECT id FROM tbl_job_weightage WHERE job_id = %s AND parameter = %s",
                (job_id, parameter)
            )
            weightage_row = process_seeker_cursor.fetchone()
            job_weightage_id = weightage_row['id'] if weightage_row else None

            process_seeker_cursor.execute("""
                INSERT INTO tbl_job_application_weightage (
                    job_application_id, job_weightage_id, parameter,
                    weightage, weightage_percent, created_by,
                    created_at, updated_by, updated_at
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                application_id, job_weightage_id, parameter,
                (divided_percent * total_parameters) / 10, divided_percent,
                1, int(time.time()), 1, int(time.time())
            ))

        process_seeker_conn.commit()

        return {
            "user_id": user_id,
            "overall_score": overall_score,
            "skills_match": match_result.get("skills_score", 0),
            "experience_match": match_result.get("experience_score", 0),
            "job_id": job_id
        }

    except Exception as e:
        return {"error": str(e)}

    finally:
        if process_seeker_cursor:
            process_seeker_cursor.close()
        if process_seeker_conn:
            process_seeker_conn.close()

def run_scoring_for_pair(user_id: int, job_id: int):
    run_conn = None
    run_cursor = None

    try:
        run_conn = connection_pool.get_connection()
        run_cursor=run_conn.cursor(dictionary=True)

        
        run_cursor.execute("""
                SELECT resume AS file_name
                FROM tbl_job_application_resume
                WHERE user_id = %s 
                AND job_id = %s 
                AND resume_current_status = 1
                ORDER BY id DESC
                LIMIT 1;

            """, (user_id, job_id))
        rr = run_cursor.fetchone()

        if not rr or not rr.get("file_name"):
                # fallback: seeker default resume
            run_cursor.execute("SELECT resume_file_name AS file_name FROM tbl_job_seeker_resume WHERE user_id = %s", (user_id,))
            rr = run_cursor.fetchone()

        if not rr or not rr.get("file_name"):
            return {"error": "No resume on file for this user/job"}

        resume_path = os.path.join(RESUMES_FILE, rr["file_name"])
        if not os.path.exists(resume_path):
            return {"error": f"Resume file missing on disk: {resume_path}"}

        resume_text = extract_text_from_file(resume_path)

        run_cursor.execute("SELECT current_ctc, year_of_experience FROM tbl_job_seeker WHERE user_id = %s", (user_id,))
        seeker = run_cursor.fetchone() or {}
        enriched_resume_text = f"""
            Resume Text:
            {resume_text}

            Seeker Details:
            - Current CTC: {seeker.get('current_ctc')}
            - Years of Experience: {seeker.get('year_of_experience')}
            """

        run_cursor.execute("SELECT description_plain_text FROM tbl_job WHERE id = %s", (job_id,))
        jr = run_cursor.fetchone()
        if not jr or not jr.get("description_plain_text"):
            return {"error": "Job description not found for given job_id"}

        run_cursor.execute("SELECT * FROM tbl_job_weightage WHERE job_id = %s", (job_id,))
        weightage_rows = run_cursor.fetchall()
        if not weightage_rows:
            return {"error": f"No weightage defined for job_id: {job_id}"}
        
        
        run_cursor.execute(
                "SELECT id FROM tbl_job_application WHERE job_id = %s AND user_id = %s",
                (job_id, user_id)
            )
        old_apps = run_cursor.fetchall()

        if old_apps:
            for row in old_apps:  # each row is a dict
                app_id = row["id"]
                    

                
                run_cursor.execute(
                    "DELETE FROM tbl_job_application_weightage WHERE job_application_id = %s",
                    (app_id,)
                )
                run_cursor.execute(
                    "DELETE FROM tbl_job_application WHERE id = %s",
                    (app_id,)
                )

                # Commit once after all deletions
            run_conn.commit()

        return process_seeker(
                job_id=job_id,
                user_id=user_id,
                resume_text=enriched_resume_text,
                weightage_rows=weightage_rows,
                job_description=jr["description_plain_text"],
            )
    except Exception as e:
        return {"error": str(e)}
    finally:
        if run_cursor: run_cursor.close()
        if run_conn: run_conn.close()


def auto_apply_function(conn,user_id,auto_apply_score):
    with conn.cursor() as cost_per_apply:
            cost_per_apply.execute('''SELECT cost_per_response 
                                   FROM tbl_master_package WHERE CODE='AAJP' ''')
            per_cost_apply=cost_per_apply.fetchone()
    auto_apply_cost=per_cost_apply['cost_per_response']



#for HR ai support
@app.route('/analyse_jobseeker', methods=['POST'])
def process_job():
    results = []
    errors = []
    start_time = time.time()
 
    job_id = request.json.get('job_id')
    if not job_id:
        #errors.append({"error": "Missing job_id"})
        return jsonify({"error": "Missing job_id"}), 400
 
    conn = None
    cursor = None
 
    try:
        conn = connection_pool.get_connection()
        cursor = conn.cursor(dictionary=True)

        cursor.execute("""
            SELECT role_id, minimum_resume_qualifying_percentage, description_plain_text,created_by
            FROM tbl_job
            WHERE id = %s
            AND CURDATE() BETWEEN post_valid_from AND post_valid_to
        """, (job_id,))
        job = cursor.fetchone()
 
        if not job:
            errors.append({"error": "Invalid job_id or job is not active"})
            cursor.execute('''UPDATE tbl_job SET ai_status = 3 WHERE id = %s''', ( job_id,))
            conn.commit()
 
            return jsonify({"error": "Invalid job_id or job is not active"}), 404
        
 
        cursor.execute("UPDATE tbl_job SET total_applications = 0,qualified_applications = 0,ai_status = 1 WHERE id = %s", (job_id,))
        conn.commit()
 
        role_id = job['role_id']
        minimum_score = job['minimum_resume_qualifying_percentage'] or 0
        description_text = job['description_plain_text']
        created_by=job['created_by']
        
        cursor.execute('''SELECT total_credits FROM tbl_user_account WHERE user_id=%s ''',(created_by,))
        user_credit_amount=cursor.fetchone()

        user_credit_amu=user_credit_amount["total_credits"]
                
        if user_credit_amu<=0:  
            errors.append(f"insufficient credit score")
            cursor.execute('''UPDATE tbl_job SET ai_status = 3 WHERE id = %s''', ( job_id,))
            conn.commit()
            return ("error:credit amount is lesser ")
        
        cursor.execute('''SELECT user_id FROM tbl_job_seeker_career_preference_roles WHERE job_role_id=%s''',(role_id,))
        preferred_roles=cursor.fetchall()  

        seeker_ids = list({r['user_id'] for r in preferred_roles if r['user_id'] is not None})
      
        if seeker_ids:
            # Create placeholders for parameterized query
            placeholders = ','.join(['%s'] * len(seeker_ids))

            query = f"""
                SELECT 
                    tjs.*,
                    scp.preferred_annual_salary,
                    GROUP_CONCAT(DISTINCT mpwl.name SEPARATOR ', ') AS preferred_locations
                FROM tbl_job_seeker tjs
                INNER JOIN tbl_users AS usr ON usr.id = tjs.user_id
                LEFT JOIN tbl_job_seeker_career_preference_locations scpl 
                    ON tjs.user_id = scpl.user_id
                LEFT JOIN tbl_job_seeker_career_preference scp 
                    ON tjs.user_id = scp.user_id
                LEFT JOIN tbl_master_preferred_work_location mpwl 
                    ON scpl.location_id = mpwl.id
                WHERE tjs.user_id IN ({placeholders})
                GROUP BY tjs.user_id;
            """
            
            cursor.execute(query, seeker_ids)
            seekers = cursor.fetchall()
        else:
            seekers = []  # no seekers found
       
        if not seekers:
            errors.append({"error": "Job seekers not found"})
            cursor.execute('''UPDATE tbl_job SET ai_status = 3 WHERE id = %s''', ( job_id,))
            conn.commit()
 
            return jsonify({"message": "Job seekers not found", "errors": errors}), 404
        
        cursor.execute('''SELECT cost_per_response FROM tbl_master_package WHERE code='CPRAIE' ''')
        score_amount=cursor.fetchone()
        print(score_amount)

        score_amu=score_amount["cost_per_response"]
        processed_seekers = []
        skipped_seekers = []

        cursor.execute('''
            INSERT INTO tbl_user_credit_debit_usage
            (user_id, user_type, usage_type, used_on,type)
            VALUES (%s, 1, 8, UNIX_TIMESTAMP(),2)
        ''', (created_by,))
        
        user_credit_usage_id = cursor.lastrowid
        
       

        for seeker in seekers:
            if user_credit_amu >= score_amu:
                current_time = int(time.time())
                        # Deduct credits
                user_credit_amu -= score_amu
                cursor.execute('''
                    INSERT INTO tbl_user_credit_debit_usage_activity
                    (user_credit_usage_id, user_id, job_id, credits_used, balance_credits, used_on)
                    VALUES (%s, %s, %s, %s, %s, %s)
                ''',(user_credit_usage_id,created_by,job_id,score_amu,user_credit_amu,current_time))
                processed_seekers.append(seeker)
            else:
                        # Not enough credits → skip rest
                skipped_seekers.append(seeker['user_id'])



                # Update DB once with total deduction
        total_deduction = len(processed_seekers) * score_amu
        remaining_balance = user_credit_amu


        cursor.execute('''
            UPDATE tbl_user_credit_debit_usage
            SET credits = %s, balance_credits = %s
            WHERE id = %s
        ''', (total_deduction, remaining_balance, user_credit_usage_id))



        cursor.execute('''
                    UPDATE tbl_user_account
                    SET total_credits = %s
                    WHERE user_id = %s
                ''', (remaining_balance, created_by))



        conn.commit()

        
        cursor.execute(
            "SELECT id FROM tbl_job_application WHERE job_id = %s",
            (job_id,)
        )
        old_apps = cursor.fetchall()
 
        if old_apps:
            for row in old_apps:  # each row is a dict
                app_id = row["id"]
               
                
                cursor.execute(
                    "DELETE FROM tbl_job_application_weightage WHERE job_application_id = %s",
                    (app_id,)
                )
                cursor.execute(
                    "DELETE FROM tbl_job_application WHERE id = %s",
                    (app_id,)
                )
 
            # Commit once after all deletions
            conn.commit()
 
        def process_single_seeker(seeker):
            thread_conn = None
            thread_cursor = None
            try:
                
                thread_conn = connection_pool.get_connection()
                thread_cursor = thread_conn.cursor(dictionary=True)
                
                
                thread_cursor.execute(
                    "SELECT * FROM tbl_job_seeker_resume WHERE user_id=%s",
                    (seeker["user_id"],)
                )
                seeker_res = thread_cursor.fetchone()
                   
                resume_path = os.path.join(RESUMES_FILE, seeker_res['resume_file_name'])
                #return resume_path
                if not os.path.exists(resume_path):
                    errors.append({"error": f"Missing resume: {seeker['resume_file_name']}"})
                    return {"error": f"Missing resume: {seeker['resume_file_name']}"}
 
                resume_text = extract_text_from_file(resume_path)
               
                enriched_resume_text = f"""
                Resume Text:
                {resume_text}
 
                Seeker Details:
                - Current CTC: {seeker['current_ctc']} 
                - Years of Experience: {seeker['year_of_experience']}
                - preferred_salary:{seeker['preferred_annual_salary']}
                - preferred_locations:{seeker['preferred_locations']}
                """
                
                
                thread_cursor.execute("SELECT * FROM tbl_job_weightage WHERE job_id = %s", (job_id,))
                weightage_rows = thread_cursor.fetchall()
               
                
                if not weightage_rows:
                    errors.append({"error": f"No weightage defined for job_id: {job_id}"})
                    return {"error": f"No weightage defined for job_id: {job_id}"}
               
                result = process_seeker(
                    job_id,
                    seeker['user_id'],
                    enriched_resume_text,
                    weightage_rows,
                    description_text
                )
 
                thread_conn.commit()
                return result
 
            except Exception as e:
                return {"error": str(e)}
 
            finally:
               if thread_cursor:
                   thread_cursor.close()
               if thread_conn:
                    thread_conn.close()
               
        with ThreadPoolExecutor(max_workers=3) as executor:
            futures = [executor.submit(process_single_seeker, seeker) for seeker in processed_seekers]
           
            for future in as_completed(futures):
                outcome = future.result()
                if outcome:
                    if "error" in outcome:
                        errors.append(outcome["error"])
                    else:
                        results.append(outcome)
 
        total_applications = len(results)
        qualified_applications = sum(1 for r in results if r.get('overall_score', 0) >= minimum_score)
       
       
        cursor.execute("""
        UPDATE tbl_job
        SET total_applications = %s,
            qualified_applications = %s,
            ai_status = 3
        WHERE id = %s
        """, (total_applications, qualified_applications, job_id))
        conn.commit()   
 
    except Exception as e:
        errors.append(str(e))
 
    finally:
        end_time = time.time()
        if conn:
            try:
                insert_ai_log(
                    conn=conn,
                    job_id=job_id,
                    user_id=None,
                    activity_type=3,
                    service_name="mistral",
                    model_name="pixtral-12b-2409",
                    user_input=None,
                    ai_response_text=json.dumps(results) if results else None,
                    response_time_ms=int((end_time - start_time) * 1000),
                    status=1 if not errors else 2,
                    error_message="; ".join([str(e) for e in errors]) if errors else None,
                    metadata=None
                )
            except Exception as log_err:
                return(f"AI Log Insertion Error: {log_err}")
            finally:
                if cursor:
                    cursor.close()
                if conn:
                    conn.close()
 
        return jsonify({"message": "Job processed.", "results": results, "errors": errors}), 200

#for seekers ai support
@app.route('/analyse_recruiter', methods=['POST'])
def init_from_seeker():
    results = []
    errors = []
    start_time = time.time()

    seeker_id = request.json.get('seeker_id')
    if not seeker_id:
        return jsonify({"error": "Missing seeker_id"}), 400

    conn = None
    cursor = None

    try:
        conn = connection_pool.get_connection()
        cursor = conn.cursor(dictionary=True)

        query = """
        SELECT 'success' AS package_status
        FROM tbl_job_seeker_package
        WHERE user_id = %s AND package_id = 6
        LIMIT 1
        """

        cursor.execute(query, (seeker_id,))
        result_status = cursor.fetchone()

        if result_status == None:
            errors.append(f"User didnt chose scoring")
            return ("error:User didnt chose scoring ")

        cursor.execute('''SELECT total_credits FROM tbl_user_account WHERE user_id=%s ''',(seeker_id,))
        user_credit_amount=cursor.fetchone()

        user_credit_amu=user_credit_amount["total_credits"]
                
        if user_credit_amu<=0:  
            errors.append(f"insufficient credit score")
            return ("error:credit amount is lesser ")
    
        cursor.execute('''SELECT 
                            tjs.*,
                            scp.preferred_annual_salary,
                            GROUP_CONCAT(DISTINCT mpwl.name SEPARATOR ', ') AS preferred_locations
                        FROM tbl_job_seeker tjs
                        LEFT JOIN tbl_job_seeker_career_preference_locations scpl 
                            ON tjs.user_id = scpl.user_id
                        LEFT JOIN tbl_job_seeker_career_preference scp 
                            ON tjs.user_id = scp.user_id
                        LEFT JOIN tbl_master_preferred_work_location mpwl 
                            ON scpl.location_id = mpwl.id
                        WHERE tjs.user_id = %s
                        GROUP BY tjs.user_id;
                        ''',(seeker_id,))
        seeker=cursor.fetchone()
       
        if not seeker:
            errors.append("Seeker not found")
            return jsonify({"message": "Seeker not found.", "results": results, "errors": errors}), 404

        cursor.execute('''SELECT * FROM tbl_job_seeker_resume WHERE user_id=%s''',(seeker_id,))
        seeker_res=cursor.fetchone()

        if not seeker_res:
            errors.append(f"Resume file not found: {seeker['resume_file_name']}")
            return jsonify({"message": "Resume not found.", "errors": ["Resume missing."]}), 404
  
        resume_path = os.path.join(RESUMES_FILE, seeker_res['resume_file_name'])
        
        if not os.path.exists(resume_path):
            errors.append(f"Resume file not found: {seeker['resume_file_name']}")
            return jsonify({"message": "Resume missing.", "results": results, "errors": errors}), 404
        
        resume_text = extract_text_from_file(resume_path)

        enriched_resume_text = f"""
        Resume Text:
        {resume_text}

        Seeker Details:
        - Current CTC: {seeker['current_ctc']}
        - Years of Experience: {seeker['year_of_experience']}
        - preferred_salary:{seeker['preferred_annual_salary']}
        - preferred_locations:{seeker['preferred_locations']}
        """
        
        cursor.execute('''SELECT job_role_id FROM tbl_job_seeker_career_preference_roles WHERE user_id=%s''',(seeker_id,))
        preferred_roles=cursor.fetchall()

        role_ids = [r['job_role_id'] for r in preferred_roles] 
        

        if role_ids:
            # create placeholders for parameterized query
            placeholders = ','.join(['%s'] * len(role_ids))

            query = f"""
                SELECT id, minimum_resume_qualifying_percentage, description_plain_text
                FROM tbl_job
                WHERE role_id IN ({placeholders})
                AND CURDATE() BETWEEN post_valid_from AND post_valid_to
            """
            
            cursor.execute(query, role_ids)
            jobs = cursor.fetchall()
        else:
            jobs = []  # no preferred roles

        
        if not jobs:
            errors.append("No active jobs found for this seeker")
            return jsonify({"message": "No active jobs found.", "results": results, "errors": errors}), 404
        
        cursor.execute('''SELECT cost_per_response FROM tbl_master_package WHERE code='GRJBR' ''')
        score_amount=cursor.fetchone()
        print(score_amount)

        score_amu=score_amount["cost_per_response"]
        processed_jobs = []
        skipped_jobs = []

        cursor.execute('''
            INSERT INTO tbl_user_credit_debit_usage
            (user_id, user_type, usage_type, used_on,type)
            VALUES (%s, 2, 6, UNIX_TIMESTAMP(),2)
        ''', (seeker_id,))
        
        user_credit_usage_id = cursor.lastrowid
       

        for job in jobs:
            if user_credit_amu >= score_amu:
                current_time = int(time.time())
                jobb_id=job['id']
                print(jobb_id)
                # Deduct credits
                user_credit_amu -= score_amu
                cursor.execute('''
                    INSERT INTO tbl_user_credit_debit_usage_activity
                    (user_credit_usage_id, user_id, job_id, credits_used, balance_credits, used_on)
                    VALUES (%s, %s, %s, %s, %s, %s)
                ''',(user_credit_usage_id,seeker_id,jobb_id,score_amu,user_credit_amu,current_time))

                # Run thread for this job
                processed_jobs.append(job)
            else:
                # Not enough credits → skip rest
                skipped_jobs.append(job['id'])

        # Update DB once with total deduction
        total_deduction = len(processed_jobs) * score_amu
        remaining_balance = user_credit_amu
  

        

        cursor.execute('''
            UPDATE tbl_user_credit_debit_usage
            SET credits = %s, balance_credits = %s
            WHERE id = %s
        ''', (total_deduction, remaining_balance, user_credit_usage_id))
      

        cursor.execute('''
            UPDATE tbl_user_account
            SET total_credits = %s
            WHERE user_id = %s
        ''', (remaining_balance, seeker_id))

        conn.commit()
       
        
        cursor.execute(
            "SELECT id FROM tbl_job_application WHERE user_id = %s",
            (seeker_id,)
        )
        old_apps = cursor.fetchall()
 
        if old_apps:
            for row in old_apps:  # each row is a dict
                app_id = row["id"]
                
                cursor.execute(
                    "DELETE FROM tbl_job_application_weightage WHERE job_application_id = %s",
                    (app_id,)
                )
                cursor.execute(
                    "DELETE FROM tbl_job_application WHERE id = %s",
                    (app_id,)
                )
 
            # Commit once after all deletions
            conn.commit()
        
        def process_job_thread(job):
            thread_conn = None
            thread_cursor = None
            try:
                thread_conn = connection_pool.get_connection()
                thread_cursor = thread_conn.cursor(dictionary=True)

                thread_cursor.execute("SELECT * FROM tbl_job_weightage WHERE job_id = %s", (job['id'],))
                weightage_rows = thread_cursor.fetchall()
            
                if not weightage_rows:
                    return {"error": f"No weightage defined for job_id: {job['id']}"}
                
                result = process_seeker(
                    job['id'],
                    seeker['user_id'],
                    enriched_resume_text, 
                    weightage_rows,
                    job['description_plain_text']
                )

                if result:
                    if result['overall_score'] >= job['minimum_resume_qualifying_percentage']:
                        thread_cursor.execute("""
                            UPDATE tbl_job SET
                                total_applications = total_applications + 1,
                                qualified_applications = qualified_applications + 1
                            WHERE id = %s
                        """, (job['id'],))
                        
                    else:
                        thread_cursor.execute("""
                            UPDATE tbl_job SET
                                total_applications = total_applications + 1
                            WHERE id = %s
                        """, (job['id'],))

                    thread_conn.commit()
                               
                return result

            except Exception as e:
                return {"error": str(e)}

            finally:
                if thread_cursor:
                    thread_cursor.close()
                if thread_conn:
                    thread_conn.close()

        #runs data asynchronous and saves additional time by running 3 data alternetively and storing output
        with ThreadPoolExecutor(max_workers=3) as executor:
            futures = [executor.submit(process_job_thread, job) for job in processed_jobs]
            for future in as_completed(futures):
                outcome = future.result()
                if outcome:
                    if "error" in outcome:
                        errors.append(outcome["error"])
                    else:
                        results.append(outcome)

                # Step 2: Fetch settings for given user_id
        cursor.execute("""
            SELECT auto_apply_score_required, min_qualification_score
            FROM tbl_job_seeker_settings
            WHERE user_id = %s
        """, (seeker_id,))
        settings = cursor.fetchone()

        if not settings:
            return jsonify({"status": 0, "message": "No settings found for this user"}), 404

        # Step 3: Check if auto_apply_score_required = 1
        if settings["auto_apply_score_required"] != 1:
            return jsonify({"status": 0, "message": "Auto apply disabled for this user"}), 200

        # Step 4: Get min_qualification_score
        min_score = settings["min_qualification_score"]

        # Step 5: Fetch job applications with overall_percent >= min_score
        cursor.execute("""
            SELECT id, job_id, overall_percent
            FROM tbl_job_application
            WHERE user_id = %s AND overall_percent >= %s
        """, (seeker_id, min_score))

        matching_jobs = cursor.fetchall()

        if not matching_jobs:
            errors.append({"status": 0, "message": "No jobs matched with required score"}), 200
        
        cursor.execute('''SELECT cost_per_response FROM tbl_master_package WHERE code='AAJP' ''')
        score_amounts=cursor.fetchone()
        print(score_amounts)

        score_amuss=score_amounts["cost_per_response"]
        autoapply_jobs = []
        skipped_auto_jobs = []

        cursor.execute('''
            INSERT INTO tbl_user_credit_debit_usage
            (user_id, user_type, usage_type, used_on,type)
            VALUES (%s, 2, 1, UNIX_TIMESTAMP(),2)
        ''', (seeker_id,))
        
        autoapply_credit_usage_id = cursor.lastrowid
        

        for job in matching_jobs:
            if remaining_balance >= score_amuss:
                current_time = int(time.time())
                remaining_balance -= score_amuss
                cursor.execute('''
                    INSERT INTO tbl_user_credit_debit_usage_activity
                    (user_credit_usage_id, user_id, job_id, credits_used, balance_credits, used_on)
                    VALUES (%s, %s, %s, %s, %s, %s)
                ''',(autoapply_credit_usage_id,seeker_id,job['job_id'],score_amu,remaining_balance,current_time))
                # Deduct credits
                
                cursor.execute("""
                   INSERT INTO tbl_job_user_action (user_id, job_id, job_application_id, action_type, action_at)
                   VALUES (%s, %s, %s, %s, %s)""",(seeker_id, job['job_id'], job['id'], "applied", current_time))


                # Run thread for this job
                autoapply_jobs.append(job)
            else:
                # Not enough credits → skip rest
                skipped_auto_jobs.append(job['id'])

        # Update DB once with total deduction
        total_deductions = len(autoapply_jobs) * score_amuss
        remaining_balance_autoapply = remaining_balance
    
        cursor.execute('''
            UPDATE tbl_user_credit_debit_usage
            SET credits = %s, balance_credits = %s
            WHERE id = %s
        ''', (total_deductions, remaining_balance_autoapply, autoapply_credit_usage_id))

        cursor.execute('''
            UPDATE tbl_user_account
            SET total_credits = %s
            WHERE user_id = %s
        ''', (remaining_balance_autoapply, seeker_id))

        conn.commit()

    except Exception as e:
        errors.append(str(e))

    finally:
        end_time = time.time()
        if conn:
            try:
                insert_ai_log(
                    conn=conn,
                    job_id=None,
                    user_id=seeker_id,
                    activity_type=6,
                    service_name="mistral",
                    model_name="pixtral-12b-2409",
                    user_input=None,
                    ai_response_text=json.dumps(results) if results else None,
                    response_time_ms=int((end_time - start_time) * 1000),
                    status=1 if not errors else 2,
                    error_message="; ".join([str(e) for e in errors]) if errors else None,
                    metadata=None
                )
            except Exception as log_err:
                return(f"AI Log Insertion Error: {log_err}")
            finally:
                if cursor:
                    cursor.close()
                if conn:
                    conn.close()

        return jsonify({"message": "Seeker processed.", "results": results, "errors": errors}), 200
 

#5 HR MAIL
SMTP_CONFIG = {
    "host": "smtp.gmail.com",
    "port": 587,
    "user": "karna1492002@gmail.com",
    "password": "lsetraeuvnafevzp"  # Gmail app password
}

OUTPUT_FOLDER = "generated_resumes"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def send_email(to_email, subject, body, attachment_path=None):
    msg = MIMEMultipart()
    msg["From"] = SMTP_CONFIG["user"]
    msg["To"] = to_email
    msg["Subject"] = subject

    msg.attach(MIMEText(body, "plain"))

    if attachment_path:
        with open(attachment_path, "rb") as f:
            attachment = MIMEApplication(f.read(), _subtype="html")
            attachment.add_header("Content-Disposition", "attachment", filename=os.path.basename(attachment_path))
            msg.attach(attachment)

    with smtplib.SMTP(SMTP_CONFIG["host"], SMTP_CONFIG["port"], timeout=30) as server:
        server.starttls()
        server.login(SMTP_CONFIG["user"], SMTP_CONFIG["password"])
        server.send_message(msg)

    print(f"📧 Email sent to {to_email}")


@app.route("/generate_job", methods=["POST"])
def generate_job():
    data = request.get_json()
    job_id = data.get("job_id")

    if not job_id:
        return jsonify({"error": "job_id is required"}), 400

    db = mysql.connector.connect(**db_config)
    cursor = db.cursor(dictionary=True)

    # --- Step 1: Check tbl_job_other_platforms ---
    cursor.execute("""
        SELECT other_platform_id FROM tbl_job_other_platforms WHERE job_id=%s
    """, (job_id,))
    platforms = [row["other_platform_id"] for row in cursor.fetchall()]

    if not platforms:
        db.close()
        return jsonify({"message": f"No platform mapping found for job_id {job_id}"}), 200

    # Only continue if platform 1 or 2 exists
    if not any(pid in (1, 2) for pid in platforms):
        db.close()
        return jsonify({"message": f"Job {job_id} not marked for LinkedIn/Career page"}), 200

    # --- Step 2: Fetch Job Details ---
    cursor.execute("""
        SELECT name, description_plain_text, about_company, cost_to_company_offer, line_manager_email_id
        FROM tbl_job WHERE id=%s
    """, (job_id,))
    job = cursor.fetchone()
    db.close()

    if not job:
        return jsonify({"error": f"No job found with id {job_id}"}), 404

    # --- Step 3: Render HTML template ---
    with open("job_template.html") as f:
        template = Template(f.read())

    html_content = template.render(
    job_id=job_id,   # ✅ pass job_id to template
    name=job["name"],
    description_plain_text=job["description_plain_text"],
    about_company=job["about_company"],
    cost_to_company_offer=job["cost_to_company_offer"],
    line_manager_email_id=job["line_manager_email_id"]
)

    # --- Save HTML file ---
    safe_name = job["name"].replace(" ", "-").lower()
    file_name = f"{job_id}-{safe_name}.html"
    file_path = os.path.join(OUTPUT_FOLDER, file_name)

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(html_content)

    print(f"✅ Job page created: {file_path}")

    # --- Send Email ---
    body = f"""
    Dear HR,

    A new job posting has been generated.

   
    Job Title: {job['name']}
    CTC: {job['cost_to_company_offer']}

    The HTML file is attached. Please use it for posting on LinkedIn and the Career Page.

    Regards,
    Job Posting System
    """

    send_email(
        to_email=job["line_manager_email_id"],
        subject=f"New Job Posting - {job['name']}",
        body=body,
        attachment_path=file_path
    )

    return jsonify({
        "message": "Job HTML created and email sent successfully",
        "file_path": file_path,
        "platforms": platforms,
        "email_sent_to": job["line_manager_email_id"]
    })


#6 voice to text
ASSEMBLYAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")
VIDEOS_FOLDER = os.path.join(os.getcwd(), "generated_resumes")
TRANSCRIPTS_FOLDER = os.path.join(os.getcwd(), "generated_resumes")
os.makedirs(TRANSCRIPTS_FOLDER, exist_ok=True)

def get_db_connection():
    return mysql.connector.connect(**db_config)

def upload_to_assemblyai(filepath):
    headers = {'authorization': ASSEMBLYAI_API_KEY}
    with open(filepath, 'rb') as f:
        response = requests.post('https://api.assemblyai.com/v2/upload', headers=headers, data=f)
    response.raise_for_status()
    return response.json()['upload_url']

def transcribe(file_url, language_code=None):
    endpoint = "https://api.assemblyai.com/v2/transcript"
    headers = {"authorization": ASSEMBLYAI_API_KEY, "content-type": "application/json"}

    if language_code:
        json_data = {
            "audio_url": file_url,
            "speaker_labels": True,
            "language_code": language_code
        }
    else:
        json_data = {
            "audio_url": file_url,
            "speaker_labels": True,
            "language_detection": True
        }

    res = requests.post(endpoint, json=json_data, headers=headers)
    res.raise_for_status()
    transcript_id = res.json()['id']

    polling_url = f"{endpoint}/{transcript_id}"
    while True:
        poll = requests.get(polling_url, headers=headers).json()
        if poll['status'] == 'completed':
            return poll['text']
        elif poll['status'] == 'error':
            raise Exception("Transcription Error", poll['error'])

def process_video(video_filename, transcript_column, job_seeker_id, cursor):
    video_path = os.path.join(VIDEOS_FOLDER, video_filename)
    if not os.path.exists(video_path):
        raise FileNotFoundError(f"Video file '{video_filename}' not found in resumes folder")

    transcript_file_path = os.path.join(
        TRANSCRIPTS_FOLDER, f"{os.path.splitext(video_filename)[0]}.txt"
    )
    file_url = upload_to_assemblyai(video_path)
    transcript = transcribe(file_url)
    with open(transcript_file_path, "w", encoding="utf-8") as f:
        f.write(transcript)
    cursor.execute(
        f"UPDATE tbl_job_seeker SET {transcript_column} = %s WHERE id = %s",
        (transcript_file_path, job_seeker_id),
    )

ASSEMBLYAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")
if not ASSEMBLYAI_API_KEY:
    raise RuntimeError("ASSEMBLYAI_API_KEY is not set in the environment.")
 
 
app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024
UPLOAD_DIR = os.path.join(os.getcwd(), "generated_resumes")
os.makedirs(UPLOAD_DIR, exist_ok=True)
 
ASSEMBLY_BASE = "https://api.assemblyai.com/v2"
HEADERS_JSON = {"authorization": ASSEMBLYAI_API_KEY, "content-type": "application/json"}
HEADERS_UPLOAD = {"authorization": ASSEMBLYAI_API_KEY}
 
def upload_to_assemblyai(local_path: str) -> str:
    with open(local_path, "rb") as f:
        resp = requests.post(f"{ASSEMBLY_BASE}/upload", headers=HEADERS_UPLOAD, data=f)
    resp.raise_for_status()
    return resp.json()["upload_url"]
 
def start_transcription(file_url: str, language_code: str | None = None) -> str:
    payload = {
        "audio_url": file_url,
        "speaker_labels": True,
    }
    if language_code:
        payload["language_code"] = language_code
    else:
        payload["language_detection"] = True
 
    resp = requests.post(f"{ASSEMBLY_BASE}/transcript", json=payload, headers=HEADERS_JSON)
    resp.raise_for_status()
    return resp.json()["id"]
 
def poll_transcription(transcript_id: str, timeout_s: int = 900, poll_interval_s: float = 3.0) -> dict:
    url = f"{ASSEMBLY_BASE}/transcript/{transcript_id}"
    start = time.time()
    while True:
        resp = requests.get(url, headers=HEADERS_JSON)
        resp.raise_for_status()
        data = resp.json()
        status = data.get("status")
 
        if status == "completed":
            return data
        if status == "error":
            raise RuntimeError(f"Transcription error: {data.get('error')}")
 
        if time.time() - start > timeout_s:
            raise TimeoutError(f"Transcription timed out after {timeout_s} seconds (id={transcript_id}).")
 
        time.sleep(poll_interval_s)
 
def transcribe_local_file(local_path: str, language_code: str | None = None) -> dict:
    upload_url = upload_to_assemblyai(local_path)
    tid = start_transcription(upload_url, language_code=language_code)
    result = poll_transcription(tid)
    return result
 
def save_upload(file_storage, field_name: str) -> str:
    filename = secure_filename(file_storage.filename or f"{field_name}.bin")
    path = os.path.join(UPLOAD_DIR, filename)
    file_storage.save(path)
    return path

# 7 build resume


def _s(v):
    if v is None:
        return ""
    if isinstance(v, (int, float)):
        return str(v)
    v = str(v).strip()
    return "" if v.lower() == "none" else v

def _allowed_template(name):
    allowed = {"modern_template.docx", "classic_template.docx", "creative_template.docx"}
    return name in allowed

def _clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"_(.*?)_", r"\1", text)
    return text.strip()

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if not underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "none")
        rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def convert_to_pdf(docx_path, output_dir):
    for cmd in ["soffice", "libreoffice"]:
        try:
            subprocess.run([
                cmd, "--headless", "--convert-to", "pdf",
                "--outdir", output_dir, docx_path
            ], check=True)
            return os.path.splitext(docx_path)[0] + ".pdf"
        except FileNotFoundError:
            continue
    raise RuntimeError("LibreOffice not found. Please install it or add to PATH.")

def fetch_resume_data_from_db(user_id: int):
    query = """
    SELECT 
    u.id AS user_id,

    (
        SELECT JSON_OBJECT(
            'full_name', CONCAT_WS(' ', uinfo.first_name, uinfo.middle_name, uinfo.last_name),
            'email', uinfo.email,
            'phone', uinfo.phone
        )
        FROM tbl_users uinfo
        WHERE uinfo.id = u.id
    ) AS user_info,

    COALESCE((
        SELECT JSON_ARRAYAGG(
            JSON_OBJECT(
                'achievement_type', m.name,
                'achievement_year', a.achievement_year,
                'description', a.discribe_about_achievement
            )
        )
        FROM tbl_job_seeker_achievement a
        JOIN tbl_master_achievement_type m ON a.achievement_type_id = m.id
        WHERE a.user_id = u.id
    ), JSON_ARRAY()) AS achievements,

    COALESCE((
        SELECT JSON_ARRAYAGG(
            JSON_OBJECT(
                'certification_name', c.certification_name,
                'valid_from', c.certification_valid_from,
                'valid_to', c.certification_valid_to,
                'certified_by', c.certified_by,
                'description', c.description
            )
        )
        FROM tbl_job_seeker_certification c
        WHERE c.user_id = u.id
    ), JSON_ARRAY()) AS certifications,

    COALESCE((
        SELECT JSON_ARRAYAGG(
            JSON_OBJECT(
                'exam_name', me.name,
                'exam_score', e.exam_score
            )
        )
        FROM tbl_job_seeker_competitive_exam e
        JOIN tbl_master_competitive_exam me ON e.competitive_exam_id = me.id
        WHERE e.user_id = u.id
    ), JSON_ARRAY()) AS competitive_exams,

    COALESCE((
        SELECT JSON_ARRAYAGG(
            JSON_OBJECT(
                'education_level', el.name,
                'examination_board', eb.name,
                'medium_of_study', ms.name,
                'percentage', edu.percentage,
                'passing_year', edu.passing_year,
                'key_skills', edu.key_skills,
                'course', co.name,
                'course_type', ct.name,
                'specialization', sp.name,
                'university', un.name,
                'starting_year', edu.starting_year,
                'ending_year', edu.ending_year,
                'grading_system', gs.name
            )
        )
        FROM tbl_job_seeker_education edu
        LEFT JOIN tbl_master_education_level el ON edu.education_level_id = el.id
        LEFT JOIN tbl_master_examination_board eb ON edu.examination_board_id = eb.id
        LEFT JOIN tbl_master_medium_of_study ms ON edu.medium_of_study_id = ms.id
        LEFT JOIN tbl_master_course co ON edu.course_id = co.id
        LEFT JOIN tbl_master_course_type ct ON edu.course_type_id = ct.id
        LEFT JOIN tbl_master_specialization sp ON edu.specialization_id = sp.id
        LEFT JOIN tbl_master_university_institute un ON edu.university_id = un.id
        LEFT JOIN tbl_master_grading_system gs ON edu.grading_system_id = gs.id
        WHERE edu.user_id = u.id
    ), JSON_ARRAY()) AS education,

    COALESCE((
        SELECT JSON_ARRAYAGG(
            JSON_OBJECT(
                'company', comp.name,
                'designation', des.name,
                'start_date', emp.start_date,
                'end_date', emp.end_date,
                'description', emp.describe_what_you_did_at_work,
                'skills',
                    (
                        SELECT COALESCE(JSON_ARRAYAGG(ms.name), JSON_ARRAY())
                        FROM tbl_job_seeker_employment_key_skills eks
                        JOIN tbl_master_key_skills ms ON eks.key_skill_id = ms.id
                        WHERE eks.job_seeker_employment_id = emp.id
                    )
            )
        )
        FROM tbl_job_seeker_employment emp
        LEFT JOIN tbl_master_company comp ON emp.company_id = comp.id
        LEFT JOIN tbl_master_designation des ON emp.designation_id = des.id
        WHERE emp.user_id = u.id
    ), JSON_ARRAY()) AS employment,

    COALESCE((
        SELECT JSON_ARRAYAGG(
            JSON_OBJECT(
                'company_name', i.company_name,
                'duration_from', i.internship_duration_from,
                'duration_to', i.internship_duration_to,
                'project_url', i.project_url,
                'description', i.discribe_about_internship,
                'skills',
                    (
                        SELECT COALESCE(JSON_ARRAYAGG(ms.name), JSON_ARRAY())
                        FROM tbl_job_seeker_internship_key_skills iks
                        JOIN tbl_master_key_skills ms ON iks.key_skill_id = ms.id
                        WHERE iks.job_seeker_internship_id = i.id
                    )
            )
        )
        FROM tbl_job_seeker_internship i
        WHERE i.user_id = u.id
    ), JSON_ARRAY()) AS internships,

    COALESCE((
        SELECT JSON_ARRAYAGG(ms.name)
        FROM tbl_job_seeker_key_skills ks
        JOIN tbl_master_key_skills ms ON ks.key_skill_id = ms.id
        WHERE ks.user_id = u.id
    ), JSON_ARRAY()) AS key_skills,

    COALESCE((
        SELECT JSON_ARRAYAGG(ml.name)
        FROM tbl_job_seeker_language jl
        JOIN tbl_master_language ml ON jl.language_id = ml.id
        WHERE jl.user_id = u.id
    ), JSON_ARRAY()) AS languages,

    COALESCE((
        SELECT JSON_ARRAYAGG(
            JSON_OBJECT(
                'project_name', p.project_name,
                'description', p.discribe_about_project,
                'key_skills',
                    (
                        SELECT COALESCE(JSON_ARRAYAGG(ms.name), JSON_ARRAY())
                        FROM tbl_job_seeker_project_key_skills pks
                        JOIN tbl_master_key_skills ms ON pks.key_skill_id = ms.id
                        WHERE pks.job_seeker_project_id = p.id
                    ),
                'project_url', p.project_url,
                'duration_from', p.project_duration_from,
                'duration_to', p.project_duration_to
            )
        )
        FROM tbl_job_seeker_project p
        WHERE p.user_id = u.id
    ), JSON_ARRAY()) AS projects,

    COALESCE((
        SELECT JSON_ARRAYAGG(
            JSON_OBJECT(
                'platform', sm.name,
                'social_media_id', js.social_media_id,
                'url', js.social_media_url
            )
        )
        FROM tbl_job_seeker_social_media js
        JOIN tbl_master_social_media sm ON js.social_media_id = sm.id
        WHERE js.user_id = u.id
    ), JSON_ARRAY()) AS social_media

FROM tbl_users u
WHERE u.id = %s;
"""

    conn = mysql.connector.connect(**db_config)
    cur = conn.cursor()
    try:
        cur.execute(query, (user_id,))
        row = cur.fetchone()
        keys = []
        if cur.description:
            keys = [desc[0] for desc in cur.description]
    finally:
        cur.close()
        conn.close()

    if not row:
        raise ValueError(f"No data found for user_id={user_id}")

    sql_data = {}
    for key, value in zip(keys, row):
        if isinstance(value, str) and (value.startswith("{") or value.startswith("[")):
            try:
                sql_data[key] = json.loads(value)
            except:
                sql_data[key] = value
        else:
            sql_data[key] = value
    return sql_data

def flatten_sql_resume_data(sql_data: dict) -> dict:
    user = sql_data.get("user_info", {}) or {}
    education = sql_data.get("education", []) or []
    employment = sql_data.get("employment", []) or []
    internships = sql_data.get("internships", []) or []
    projects = sql_data.get("projects", []) or []
    certifications = sql_data.get("certifications", []) or []
    achievements = sql_data.get("achievements", []) or []
    skills = sql_data.get("key_skills", []) or []
    languages = sql_data.get("languages", []) or []

    flattened = {
        "FULL_NAME": _s(user.get("full_name")),
        "PHONE": _s(user.get("phone")),
        "EMAIL": _s(user.get("email")),
        "TECHNICAL_SKILLS": ", ".join(_s(s) for s in skills if s),
        "LANGUAGES": ", ".join(_s(l) for l in languages if l),
    }

    edu_text = [
        f"{_s(e.get('course'))} in {_s(e.get('specialization'))} - {_s(e.get('university'))} "
        f"({_s(e.get('starting_year'))}–{_s(e.get('ending_year'))})"
        for e in education
        if isinstance(e, dict) and any(_s(v) for v in e.values())
    ]
    flattened["EDUCATION_SECTION"] = "\n".join(edu_text)

    exp_text = [
        f"{_s(emp.get('designation'))} at {_s(emp.get('company'))} "
        f"({_s(emp.get('start_date'))}–{_s(emp.get('end_date'))})\n{_s(emp.get('description'))}"
        for emp in employment
        if isinstance(emp, dict) and any(_s(v) for v in emp.values())
    ]
    flattened["EXPERIENCE"] = "\n\n".join(exp_text)

    intern_text = [
        f"{_s(i.get('company_name'))} ({_s(i.get('duration_from'))}–{_s(i.get('duration_to'))})\n"
        f"{_s(i.get('description'))}\n"
        f"Skills: {', '.join(_s(s) for s in i.get('skills', []) if s)}\n"
        f"URL: {_s(i.get('project_url'))}"
        for i in internships
        if isinstance(i, dict) and any(_s(v) for v in i.values())
    ]
    flattened["INTERNSHIPS"] = "\n\n".join(intern_text)

    proj_text = [
        f"{_s(p.get('project_name'))} ({_s(p.get('duration_from'))}–{_s(p.get('duration_to'))})\n"
        f"{_s(p.get('description'))}\n"
        f"Skills: {', '.join(_s(s) for s in p.get('key_skills', []) if s)}"
        for p in projects
        if isinstance(p, dict) and any(_s(v) for v in p.values())
    ]
    flattened["PROJECTS"] = "\n\n".join(proj_text)

    cert_text = [
        f"{_s(c.get('certification_name'))} – {_s(c.get('certified_by'))} "
        f"({_s(c.get('valid_from'))}–{_s(c.get('valid_to'))})"
        for c in certifications
        if isinstance(c, dict) and any(_s(v) for v in c.values())
    ]
    flattened["CERTIFICATIONS"] = "\n".join(cert_text)

    ach_text = [
        f"{_s(a.get('achievement_type'))} ({_s(a.get('achievement_year'))}): {_s(a.get('description'))}"
        for a in achievements
        if isinstance(a, dict) and any(_s(v) for v in a.values())
    ]
    flattened["ACHIEVEMENTS"] = "\n".join(ach_text)

    flattened["SUMMARY"] = ""
    return flattened

def parse_json_response(content, fallback):
    try:
        return json.loads(content)
    except Exception:
        match = re.search(r"\{.*\}", content, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except Exception:
                pass
    return fallback or {}

def enhance_resume_all(data):
    prompt = f"""
[ROLE]
You are an experienced professional resume writer and editor.
Your job is to professionally enhance the content below while keeping all factual details accurate.

[INPUT]
Here is the candidate's résumé data in JSON format:
{json.dumps(data, indent=2)}

[STEPS]
1. Improve phrasing in each section to sound more professional and achievement-oriented.
2. Use strong action verbs and recruiter-friendly language (e.g., “developed,” “optimized,” “engineered”).
3. Preserve factual information — do NOT add or guess new details.
4. If SUMMARY is blank, create a short professional summary within 150 words based on the data.
5. Empty or junk content should become "" (empty string) or an empty array [].
Input project: "description": "qwefwef" → Do NOT rewrite as "Developed advanced systems". REMOVE it.
6. If any list such as "key_skills", "skills", or "key_skills" is empty ( [] ), 
   → do NOT create placeholder text like "Skills: [List of skills used]".
   → simply omit that line or keep the array empty.

[EXPECTATION]
Return valid JSON only, preserving the same keys and structure.
Each value should be grammatically refined, polished, and resume-appropriate.
Keys: {list(data.keys())}
"""
    try:
        response = mistral_client.chat.complete(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=1500
        )
        content = (response.choices[0].message.content or "").strip()
        parsed = parse_json_response(content, data)
        if not isinstance(parsed, dict):
            print("⚠️ Warning: Mistral returned non-dict, using fallback.")
            return data
        return parsed
    except Exception as e:
        print("⚠️ Mistral API error:", e)
        return data

def fill_template(data, output_path, template_path):
    doc = Document(template_path)
    def is_emptyish(val) -> bool:
        if val is None:
            return True
        s = str(val).strip()
        return s == "" or s.lower() in {"none", "n/a", "na", "null", "[]", "{}", "0"}

    def clear_paragraph(paragraph):
        for r in paragraph.runs:
            r.text = ""

    def insert_with_format(paragraph, text, bold_all=False, uppercase=False, font_size=None):
        if text is None:
            return
        text = str(text)
        if bold_all:
            run = paragraph.add_run(text.upper() if uppercase else text)
            run.bold = True
            if font_size:
                run.font.size = Pt(font_size)
            return
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if not part:
                continue
            clean = part.replace("**", "")
            run = paragraph.add_run(clean.upper() if uppercase else clean)
            if part.startswith("**") and part.endswith("**"):
                run.bold = True
            if font_size:
                run.font.size = Pt(font_size)

    def normalize_header_text(txt):
        return re.sub(r'\s+', ' ', txt.strip().upper())

    header_map = {
        "SUMMARY": "PROFILE",
        "EDUCATION_SECTION": "EDUCATION",
        "TECHNICAL_SKILLS": "TECHNICAL SKILLS",
        "CERTIFICATIONS": "CERTIFICATIONS",
        "EXPERIENCE": "WORK EXPERIENCE",
        "PROJECTS": "PROJECTS",
        "INTERNSHIPS": "INTERNSHIPS",
        "ACHIEVEMENTS": "ACHIEVEMENTS",
        "LANGUAGES": "LANGUAGES",
        "PHONE": "CONTACT",
        "EMAIL": "CONTACT",
    }

    headers_to_keep = {header: False for header in header_map.values()}
    for field, header in header_map.items():
        if not is_emptyish(data.get(field, "")):
            headers_to_keep[header] = True

    for para in list(doc.paragraphs):
        text = para.text
        if not text:
            continue
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in text:
                val_str = "" if value is None else str(value).strip()
                clear_paragraph(para)
                if is_emptyish(val_str):
                    continue
                if key == "FULL_NAME":
                    insert_with_format(para, val_str, bold_all=True, uppercase=True, font_size=16)
                else:
                    insert_with_format(para, val_str)
                break

    headers_to_remove = {h.upper() for h, keep in headers_to_keep.items() if not keep}

    def remove_paragraph_at_index(idx):
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)

    i = 0
    while i < len(doc.paragraphs):
        txt = normalize_header_text(doc.paragraphs[i].text)
        if any(h in txt for h in headers_to_remove):
            remove_paragraph_at_index(i)
            while i < len(doc.paragraphs) and not doc.paragraphs[i].text.strip():
                remove_paragraph_at_index(i)
            continue
        i += 1

    doc.save(output_path)

@app.route("/build-resume", methods=["POST"])
def build_resume():
    start_time = time.time()
    conn = None
    cursor = None
    try:
        data = request.get_json()
        user_id = data.get("user_id")
        if not user_id:
            return jsonify({"error": "user_id is required"}), 400

        conn = connection_pool.get_connection()
        cursor = conn.cursor(dictionary=True)

        # 1️⃣ Fetch cost_per_response for Build Resume
        cursor.execute("""
            SELECT COALESCE(cost_per_response, 0) AS cost_per_response
            FROM tbl_master_package
            WHERE id = 5
        """)
        row = cursor.fetchone()
        cost_per_response = Decimal(str(row["cost_per_response"])) if row else Decimal("0")

        # 2️⃣ Check user credits
        cursor.execute("""
            SELECT COALESCE(total_credits, 0) AS total_credits
            FROM tbl_user_account
            WHERE user_id = %s
            FOR UPDATE
        """, (user_id,))
        row = cursor.fetchone()
        if not row:
            conn.rollback()
            return jsonify({"error": "User account not found"}), 404

        total_credits = Decimal(str(row["total_credits"]))
        if total_credits < cost_per_response:
            conn.rollback()
            return jsonify({
                "status": 0,
                "error": {"message": "Insufficient credits"}
            }), 402

        # 3️⃣ Fetch data and enhance
        sql_data = fetch_resume_data_from_db(int(user_id))
        raw_data = flatten_sql_resume_data(sql_data)
        enhanced_data = enhance_resume_all(raw_data)
        if not isinstance(enhanced_data, dict):
            enhanced_data = raw_data

        # 4️⃣ Prepare template + output
        template_path = TEMPLATE_PATH
        filename = f"templated_resume_{uuid.uuid4().hex}.docx"
        output_docx = os.path.join(OUTPUT_DIR, filename)
        output_pdf = os.path.splitext(output_docx)[0] + ".pdf"

        fill_template(enhanced_data, output_docx, template_path)
        convert_to_pdf(output_docx, OUTPUT_DIR)

        # 5️⃣ Safe retry function for Windows lock issue
        def safe_build_file_format(path, retries=5, delay=0.5):
            for attempt in range(retries):
                try:
                    return build_file_format(path)
                except PermissionError:
                    if attempt == retries - 1:
                        raise
                    time.sleep(delay)

        cv_format = safe_build_file_format(output_pdf)

        # 6️⃣ Update generated resume info into tbl_job_seeker_resume
        try:
            cursor.execute("""
                UPDATE tbl_job_seeker_resume
                SET resume_file_name = %s,
                    file_data = %s
                WHERE user_id = %s
            """, (os.path.basename(output_pdf), json.dumps(cv_format, ensure_ascii=False).encode("utf-8"),user_id))
            conn.commit()
        except Exception as resume_update_err:
            print(f"⚠️ Failed to update resume info: {resume_update_err}")
            conn.rollback()

        # 7️⃣ Deduct credits and log usage
        new_balance = total_credits - cost_per_response

        cursor.execute("""
            UPDATE tbl_user_account
            SET total_credits = %s
            WHERE user_id = %s
        """, (new_balance, user_id))
        conn.commit()

        cursor.execute("""
            INSERT INTO tbl_user_credit_debit_usage
                (user_id, user_type, credits, usage_type, used_on, balance_credits, type)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (user_id, 2, cost_per_response, 2, int(time.time()), new_balance, 2))
        usage_id = cursor.lastrowid

        cursor.execute("""
            INSERT INTO tbl_user_credit_debit_usage_activity
                (user_credit_usage_id, user_id, job_id, credits_used, balance_credits)
            VALUES (%s, %s, NULL, %s, %s)
        """, (usage_id, user_id, cost_per_response, new_balance))
        conn.commit()

        # 8️⃣ Build response
        response_payload = {
            "status": "1",
            "message": "Resume built successfully",
            "data": {
                "curriculum_vitae": os.path.basename(output_pdf),
                "curriculum_vitae_format": cv_format
            }
        }

        # 9️⃣ Log AI activity
        try:
            response_time_ms = int((time.time() - start_time) * 1000)
            conn_log = connection_pool.get_connection()

            user_input = {
                "user_id": user_id,
                "template": os.path.basename(template_path),
                "request_data": {k: v for k, v in raw_data.items() if v}
            }

            ai_response = {
                "enhanced_data": enhanced_data,
                "output_files": {
                    "docx": os.path.basename(output_docx),
                    "pdf": os.path.basename(output_pdf)
                }
            }

            metadata = {
                "credits_deducted": float(cost_per_response),
                "balance_after_deduction": float(new_balance),
                "status": "success"
            }

            insert_ai_log(
                conn=conn_log,
                job_id=None,
                user_id=user_id,
                activity_type="2",  # Build Resume
                service_name="mistral",
                model_name=MODEL,
                user_input=json.dumps(user_input, ensure_ascii=False),
                ai_response_text=json.dumps(ai_response, ensure_ascii=False),
                response_time_ms=response_time_ms,
                status=1,
                error_message="",
                metadata=json.dumps(metadata, ensure_ascii=False)
            )
            conn_log.close()
        except Exception as log_err:
            print(f"⚠️ Failed to insert AI log for build resume: {log_err}")

        return jsonify(response_payload), 200

    except Exception as e:
        traceback.print_exc()
        if conn:
            conn.rollback()
        return jsonify({"error": str(e)}), 500

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()



if __name__ =='__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
